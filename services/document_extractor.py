#!/usr/bin/env python3
"""
Document Text Extraction Script
Handles PDF, Excel, and Word document text extraction for the Extractly platform.
"""

import sys
import json
import base64
import io
import tempfile
import os
from typing import List, Dict, Any

# Document processing libraries
try:
    import PyPDF2
    import pandas as pd
    from docx import Document
    import xlrd
    from openpyxl import load_workbook
    from pdfminer.high_level import extract_text as pdfminer_extract_text
    from pdfminer.pdfparser import PDFSyntaxError
except ImportError as e:
    print(f"Error: Missing required library: {e}", file=sys.stderr)
    sys.exit(1)

def merge_multirow_headers(rows):
    """
    Conservatively detect and merge multi-row headers in Excel data.
    
    Only merges if we're confident rows are headers (not data).
    Returns rows unchanged unless multiple consecutive sparse long-text 
    header rows are found starting from row 0.
    
    Merge criteria:
    - Rows must have ≤5 non-blank cells (sparse, like multi-row headers)
    - At least one cell must have long text (>25 chars)
    - Rows must be consecutive starting from row 0
    
    This prevents data rows from being merged into headers.
    """
    if not rows or len(rows) <= 1:
        return rows
    
    # Check if we even have multi-row headers
    # Look for pattern: rows with mostly blanks and long text in some columns
    header_candidate_rows = []
    
    for i, row in enumerate(rows[:5]):  # Check first 5 rows only
        non_blank_count = sum(1 for cell in row if cell != "blank")
        
        # Header rows typically have FEW non-blank cells with LONG descriptive text
        if non_blank_count <= 5:  # Very few non-blank cells
            # Check if the non-blank cells have long header-like text
            has_header_text = False
            for cell in row:
                if cell != "blank" and len(cell) > 25:  # Long descriptive text
                    has_header_text = True
                    break
            
            if has_header_text:
                header_candidate_rows.append(i)
        else:
            # Many non-blank cells = likely data row, stop looking
            break
    
    # If we found multiple consecutive header rows starting from row 0
    if len(header_candidate_rows) > 1 and header_candidate_rows[0] == 0:
        # Check if they're actually consecutive
        is_consecutive = all(
            header_candidate_rows[i] == header_candidate_rows[i-1] + 1 
            for i in range(1, len(header_candidate_rows))
        )
        
        if is_consecutive:
            header_end_index = header_candidate_rows[-1] + 1
            header_rows = rows[:header_end_index]
            data_rows = rows[header_end_index:]
            
            # Merge header rows
            merged_header = merge_header_rows(header_rows)
            return [merged_header] + data_rows
    
    # Default: return rows unchanged (no merging detected)
    return rows


def merge_header_rows(header_rows):
    """
    Merge multiple header rows into a single header row.
    
    For each column position, combine text from all header rows
    where that column has non-blank values.
    """
    if not header_rows:
        return []
    
    num_columns = max(len(row) for row in header_rows)
    merged_header = []
    
    for col_index in range(num_columns):
        # Collect all non-blank values from this column across header rows
        column_parts = []
        
        for row in header_rows:
            if col_index < len(row) and row[col_index] != "blank":
                column_parts.append(row[col_index])
        
        # Combine the parts with spaces
        if column_parts:
            merged_header.append(" ".join(column_parts))
        else:
            merged_header.append("blank")
    
    return merged_header

def extract_pdf_text_with_gemini(file_content: bytes, file_name: str = "document.pdf") -> str:
    """Use Gemini AI to extract text from a scanned/image-based PDF."""
    try:
        import google.generativeai as genai
        api_key = os.environ.get("GOOGLE_API_KEY") or os.environ.get("GEMINI_API_KEY")
        if not api_key:
            raise Exception("No Gemini API key available")
        
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-2.0-flash")
        
        pdf_b64 = base64.b64encode(file_content).decode('utf-8')
        
        response = model.generate_content([
            {
                "mime_type": "application/pdf",
                "data": pdf_b64
            },
            "Extract ALL text content from this document. Return only the raw text content, preserving the structure and formatting as closely as possible. Do not add any commentary or explanation."
        ])
        
        if response and response.text:
            return response.text.strip()
        return ""
    except Exception as e:
        print(f"Gemini PDF extraction failed: {str(e)}", file=sys.stderr)
        return ""

def extract_pdf_text(file_content: bytes, file_name: str = "document.pdf") -> str:
    """Extract text from PDF file using PyPDF2 with pdfminer fallback, then Gemini for scanned docs."""
    text = ""
    
    # Try PyPDF2 first
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        print(f"PyPDF2 extraction failed: {str(e)}", file=sys.stderr)
    
    # If PyPDF2 didn't extract much text, try pdfminer as fallback
    if len(text.strip()) < 50:
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                tmp_file.write(file_content)
                tmp_file.flush()
                pdfminer_text = pdfminer_extract_text(tmp_file.name)
                os.unlink(tmp_file.name)
                if pdfminer_text and len(pdfminer_text.strip()) > len(text.strip()):
                    text = pdfminer_text
        except Exception as e:
            print(f"pdfminer extraction failed: {str(e)}", file=sys.stderr)
    
    # If still no text, try Gemini AI for scanned/image-based PDFs
    if not text.strip():
        print(f"No text extracted with standard methods, trying Gemini AI for OCR...", file=sys.stderr)
        gemini_text = extract_pdf_text_with_gemini(file_content, file_name)
        if gemini_text:
            return gemini_text
        raise Exception("PDF extraction failed: No text could be extracted (may be scanned/image-based)")
    
    return text.strip()

def extract_docx_text(file_content: bytes) -> str:
    """Extract text from DOCX file."""
    try:
        doc = Document(io.BytesIO(file_content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        
        return text.strip()
    except Exception as e:
        raise Exception(f"DOCX extraction failed: {str(e)}")

def extract_excel_text(file_content: bytes, file_name: str) -> str:
    """Extract text from Excel file (both .xls and .xlsx) - extracts ALL rows."""
    text_parts = []
    
    # Try modern Excel format first (.xlsx)
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_file:
            tmp_file.write(file_content)
            tmp_file.flush()
            
            workbook = load_workbook(tmp_file.name, data_only=True)
            
            for sheet_name in workbook.sheetnames:
                worksheet = workbook[sheet_name]
                text_parts.append(f"=== Sheet: {sheet_name} ===")
                
                # Get the actual dimensions - use robust column detection
                max_row = worksheet.max_row
                
                # Find ACTUAL maximum column by scanning first few rows extensively
                # Don't trust worksheet.max_column as it can miss columns with formatting/empty cells
                actual_max_col = 0
                for row_num in range(1, min(15, max_row + 1)):  # Check first 15 rows
                    for col_num in range(1, 200):  # Check up to column 200 to handle wide sheets
                        cell_value = worksheet.cell(row=row_num, column=col_num).value
                        if cell_value is not None and str(cell_value).strip():
                            actual_max_col = max(actual_max_col, col_num)
                
                # Use the detected column count or minimum of 150 columns for wide sheets
                max_col = max(actual_max_col, worksheet.max_column, 150)
                
                # Smart header detection and multi-row header merging
                all_rows = []
                for row_num in range(1, max_row + 1):
                    row_data = []
                    for col_num in range(1, max_col + 1):
                        cell_value = worksheet.cell(row=row_num, column=col_num).value
                        if cell_value is not None and str(cell_value).strip():
                            row_data.append(str(cell_value))
                        else:
                            row_data.append("blank")
                    
                    # Only include non-empty rows
                    if any(cell != "blank" for cell in row_data):
                        all_rows.append(row_data)
                
                if all_rows:
                    # Detect header rows and merge multi-row headers
                    merged_rows = merge_multirow_headers(all_rows)
                    for row in merged_rows:
                        text_parts.append("\t".join(row))
            
            os.unlink(tmp_file.name)
            return "\n".join(text_parts)
            
    except Exception:
        # Fall back to older Excel format (.xls)
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.xls') as tmp_file:
                tmp_file.write(file_content)
                tmp_file.flush()
                
                workbook = xlrd.open_workbook(tmp_file.name)
                
                for sheet_index in range(workbook.nsheets):
                    sheet = workbook.sheet_by_index(sheet_index)
                    text_parts.append(f"=== Sheet: {sheet.name} ===")
                    
                    # Find ACTUAL maximum column by scanning extensively
                    # Don't trust sheet.ncols as it can miss columns with formatting/empty cells
                    actual_max_cols = 0
                    for row_index in range(min(15, sheet.nrows)):  # Check first 15 rows
                        for col_index in range(200):  # Check up to column 200 to handle wide sheets
                            try:
                                cell_value = sheet.cell_value(row_index, col_index)
                                if cell_value is not None and str(cell_value).strip():
                                    actual_max_cols = max(actual_max_cols, col_index + 1)  # +1 because col_index is 0-based
                            except:
                                break  # xlrd will throw exception when past actual columns
                    
                    # Use the detected column count or minimum of 150 columns for wide sheets
                    max_cols = max(actual_max_cols, sheet.ncols, 150)
                    
                    # Smart header detection and multi-row header merging
                    all_rows = []
                    for row_index in range(sheet.nrows):
                        row_data = []
                        for col_index in range(max_cols):
                            try:
                                cell_value = sheet.cell_value(row_index, col_index)
                                if cell_value is not None and str(cell_value).strip():
                                    row_data.append(str(cell_value))
                                else:
                                    row_data.append("blank")
                            except:
                                row_data.append("blank")
                        
                        # Only include non-empty rows
                        if any(cell != "blank" for cell in row_data):
                            all_rows.append(row_data)
                    
                    if all_rows:
                        # Detect header rows and merge multi-row headers
                        merged_rows = merge_multirow_headers(all_rows)
                        for row in merged_rows:
                            text_parts.append("\t".join(row))
                
                os.unlink(tmp_file.name)
                return "\n".join(text_parts)
                
        except Exception as xls_error:
            # Final fallback using pandas
            try:
                excel_data = pd.read_excel(io.BytesIO(file_content), sheet_name=None)
                
                for sheet_name, df in excel_data.items():
                    text_parts.append(f"=== Sheet: {sheet_name} ===")
                    
                    # Include headers with consistent column count
                    headers = df.columns.tolist()
                    text_parts.append("\t".join(str(h) for h in headers))
                    
                    # Include ALL data rows with consistent column count
                    for _, row in df.iterrows():
                        # Ensure every row has same number of columns as headers
                        row_data = []
                        for val in row.values:
                            if pd.notna(val) and str(val).strip():
                                row_data.append(str(val))
                            else:
                                row_data.append("blank")
                        
                        # Only skip completely empty rows (all blanks)
                        if any(cell != "blank" for cell in row_data):
                            text_parts.append("\t".join(row_data))
                
                return "\n".join(text_parts)
                
            except Exception as pandas_error:
                raise Exception(f"Excel extraction failed with all methods. XLS error: {str(xls_error)}, Pandas error: {str(pandas_error)}")

def extract_text_from_document(file_data: Dict[str, Any]) -> Dict[str, Any]:
    """Extract text from a single document."""
    file_name = file_data.get('file_name', '')
    file_content_b64 = file_data.get('file_content', '')
    mime_type = file_data.get('mime_type', '')
    
    # Handle data URL format (data:mime/type;base64,content)
    if file_content_b64.startswith('data:'):
        # Split off the data URL prefix
        _, file_content_b64 = file_content_b64.split(',', 1)
    
    try:
        file_content = base64.b64decode(file_content_b64)
    except Exception as e:
        return {
            'file_name': file_name,
            'text_content': '',
            'error': f'Base64 decode failed: {str(e)}',
            'file_size': 0,
            'word_count': 0,
            'extraction_method': 'failed'
        }
    
    file_size = len(file_content)
    
    try:
        # Determine extraction method based on MIME type and file extension
        file_ext = os.path.splitext(file_name.lower())[1]
        
        if mime_type == 'application/pdf' or file_ext == '.pdf':
            extracted_text = extract_pdf_text(file_content)
        elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword'] or file_ext in ['.docx', '.doc']:
            extracted_text = extract_docx_text(file_content)
        elif mime_type in ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', 'application/vnd.ms-excel'] or file_ext in ['.xlsx', '.xls']:
            extracted_text = extract_excel_text(file_content, file_name)
        else:
            return {
                'file_name': file_name,
                'text_content': '',
                'error': f'Unsupported file type: {mime_type} ({file_ext})',
                'file_size': file_size,
                'word_count': 0,
                'extraction_method': 'unsupported'
            }
        
        return {
            'file_name': file_name,
            'text_content': extracted_text,
            'file_size': file_size,
            'mime_type': mime_type,
            'word_count': len(extracted_text.split()) if extracted_text else 0,
            'extraction_method': 'direct'
        }
        
    except Exception as e:
        return {
            'file_name': file_name,
            'text_content': '',
            'error': str(e),
            'file_size': file_size,
            'mime_type': mime_type,
            'word_count': 0,
            'extraction_method': 'failed'
        }

def main():
    """Main function to process extraction requests."""
    try:
        # Read input from stdin
        input_data = sys.stdin.read()
        request = json.loads(input_data)
        
        step = request.get('step', '')
        documents = request.get('documents', [])
        
        if step not in ['extract_text_only', 'extract']:
            print(json.dumps({
                'success': False,
                'error': f'Unsupported step: {step}. Only "extract_text_only" and "extract" are supported.'
            }))
            return
        
        # Process each document
        extracted_texts = []
        for doc in documents:
            result = extract_text_from_document(doc)
            extracted_texts.append(result)
        
        # Return results
        output = {
            'success': True,
            'extracted_texts': extracted_texts
        }
        
        print(json.dumps(output))
        
    except json.JSONDecodeError as e:
        print(json.dumps({
            'success': False,
            'error': f'Invalid JSON input: {str(e)}'
        }))
    except Exception as e:
        print(json.dumps({
            'success': False,
            'error': f'Processing failed: {str(e)}'
        }), file=sys.stderr)

if __name__ == '__main__':
    main()