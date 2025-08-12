#!/usr/bin/env python3
"""
SIMPLIFIED AI EXTRACTION SYSTEM
Two-step process: Extract → Validate
Each step can be called individually or chained together
"""
import os
import json
import logging
import base64
from typing import Dict, Any, List, Optional
from dataclasses import dataclass
from prompt import EXTRACTION_PROMPT

# Configure logging
logging.basicConfig(level=logging.INFO)

@dataclass
class ExtractionResult:
    success: bool
    extracted_data: Optional[Dict[str, Any]] = None
    error_message: Optional[str] = None
    extraction_prompt: Optional[str] = None
    ai_response: Optional[str] = None
    input_token_count: Optional[int] = None
    output_token_count: Optional[int] = None

# ValidationResult dataclass removed - validation now occurs only during extraction

def process_excel_file_content(file_content: str, file_name: str) -> str:
    """Process Excel file content using existing pandas/openpyxl processing (same as used in main extraction)"""
    import base64
    import io
    
    try:
        # Import required libraries (already available from existing system)
        import pandas as pd
        import openpyxl
        from openpyxl import load_workbook
        
        # Convert base64 to bytes
        if file_content.startswith('data:'):
            file_content = file_content.split(',')[1]
        
        file_bytes = base64.b64decode(file_content)
        file_buffer = io.BytesIO(file_bytes)
        
        # Use the same approach as the existing Excel processing
        extracted_content = ["Excel file content:"]
        
        try:
            # Try openpyxl first for .xlsx files
            workbook = load_workbook(file_buffer, read_only=True, data_only=True)
            
            for sheet_name in workbook.sheetnames:
                extracted_content.append(f"=== SHEET: {sheet_name} ===")
                
                worksheet = workbook[sheet_name]
                row_count = 0
                
                # Get data with limited rows for text representation (similar to existing pattern)
                for row in worksheet.iter_rows(values_only=True):
                    if row_count > 50:  # Limit for text extraction phase
                        break
                        
                    # Convert None values to empty strings
                    cleaned_row = []
                    for cell in row:
                        if cell is None:
                            cleaned_row.append("")
                        else:
                            cleaned_row.append(str(cell))
                    
                    # Skip completely empty rows
                    if any(cell.strip() for cell in cleaned_row if cell):
                        extracted_content.append("    " + "    ".join(cleaned_row[:10]))  # First 10 columns
                        row_count += 1
                
                if row_count >= 50:
                    extracted_content.append(f"    ... (showing first 50 rows of {sheet_name})")
            
            workbook.close()
            
        except Exception as openpyxl_error:
            # Fallback to pandas (existing pattern)
            try:
                file_buffer.seek(0)
                excel_data = pd.read_excel(file_buffer, sheet_name=None, nrows=50)  # Limit rows
                
                for sheet_name, df in excel_data.items():
                    extracted_content.append(f"=== SHEET: {sheet_name} ===")
                    
                    if not df.empty:
                        # Column headers
                        headers = "    " + "    ".join(str(col) for col in df.columns[:10])
                        extracted_content.append(headers)
                        
                        # Data rows (limit to first 50)
                        for idx, (_, row) in enumerate(df.iterrows()):
                            if idx >= 49:  # 0-based, so 49 = 50th row
                                break
                            row_data = "    " + "    ".join(str(val) if pd.notna(val) else "" for val in row[:10])
                            extracted_content.append(row_data)
                        
                        if len(df) >= 50:
                            extracted_content.append(f"    ... (showing first 50 rows of {sheet_name})")
                    
            except Exception as pandas_error:
                return f"Error reading Excel file: {pandas_error}"
        
        result = "\n".join(extracted_content)
        return result if result else f"No readable data in Excel file {file_name}"
        
    except Exception as e:
        logging.error(f"Error processing Excel content from {file_name}: {e}")
        return f"Error processing Excel file: {e}"

def extract_pdf_with_gemini(file_content: str, file_name: str) -> str:
    """Extract content from PDF files using Gemini API"""
    import base64
    import os
    from google import genai
    from google.genai import types
    
    try:
        # Try GEMINI_API_KEY first, then GOOGLE_API_KEY as fallback
        api_key = os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
        if not api_key:
            return f"GEMINI_API_KEY or GOOGLE_API_KEY not found for extracting {file_name}"
        
        logging.info(f"Using API key for PDF extraction: {'GEMINI_API_KEY' if os.environ.get('GEMINI_API_KEY') else 'GOOGLE_API_KEY'}")
        client = genai.Client(api_key=api_key)
        
        # Convert base64 to bytes
        if file_content.startswith('data:'):
            file_content = file_content.split(',')[1]
        
        file_bytes = base64.b64decode(file_content)
        
        prompt = """Extract all text content from this PDF document. 
        Return the complete text content exactly as it appears, preserving formatting and structure.
        Do not summarize or modify the content - extract everything."""
        
        logging.info(f"Sending PDF {file_name} to Gemini (size: {len(file_bytes)} bytes)")
        
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=[
                types.Part.from_bytes(
                    data=file_bytes,
                    mime_type="application/pdf"
                ),
                prompt
            ],
            config=types.GenerateContentConfig(
                max_output_tokens=50000,
                temperature=0.1
            )
        )
        
        logging.info(f"Gemini response received for {file_name}")
        
        if response and response.text:
            content = response.text.strip()
            logging.info(f"Extracted {len(content)} characters from {file_name}")
            logging.info(f"First 100 chars: {content[:100]}...")
            return content
        elif response:
            logging.warning(f"Gemini response exists but no text for {file_name}")
            logging.warning(f"Response object: {str(response)[:200]}")
            return f"No content extracted from {file_name} - Gemini response had no text"
        else:
            logging.warning(f"No response from Gemini for {file_name}")
            return f"No content extracted from {file_name} - Gemini returned no response"
            
    except Exception as e:
        logging.error(f"Error extracting PDF content from {file_name}: {e}")
        logging.error(f"Exception type: {type(e).__name__}")
        return f"Error extracting PDF content: {e}"

def extract_document_with_gemini(file_content: str, file_name: str, mime_type: str) -> str:
    """Extract content from other document types using Gemini API"""
    import base64
    import os
    from google import genai
    from google.genai import types
    
    try:
        api_key = os.environ.get("GEMINI_API_KEY")
        if not api_key:
            return f"GEMINI_API_KEY not found for extracting {file_name}"
        
        client = genai.Client(api_key=api_key)
        
        # Convert base64 to bytes
        if file_content.startswith('data:'):
            file_content = file_content.split(',')[1]
        
        file_bytes = base64.b64decode(file_content)
        
        prompt = """Extract all text content from this document. 
        Return the complete text content exactly as it appears, preserving formatting and structure.
        Do not summarize or modify the content - extract everything."""
        
        # Map common file extensions to MIME types
        if not mime_type or mime_type == 'application/octet-stream':
            if file_name.lower().endswith('.docx'):
                mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            elif file_name.lower().endswith('.doc'):
                mime_type = "application/msword"
            else:
                mime_type = "application/octet-stream"
        
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=[
                types.Part.from_bytes(
                    data=file_bytes,
                    mime_type=mime_type
                ),
                prompt
            ],
            config=types.GenerateContentConfig(
                max_output_tokens=50000,
                temperature=0.1
            )
        )
        
        if response and response.text:
            return response.text.strip()
        else:
            return f"No content extracted from {file_name}"
            
    except Exception as e:
        logging.error(f"Error extracting document content from {file_name}: {e}")
        return f"Error extracting document content: {e}"

def repair_truncated_json(response_text: str) -> Optional[str]:
    """
    Enhanced repair function for truncated JSON responses. Finds complete field validation objects
    and properly closes the JSON structure to preserve as much data as possible.
    """
    try:
        logging.info(f"Attempting to repair JSON response of length {len(response_text)}")
        
        # Check if response starts with field_validations structure
        if not response_text.strip().startswith('{"field_validations":'):
            logging.warning("Response doesn't start with expected field_validations structure")
            # Try alternative patterns
            if '"field_validations"' in response_text[:200]:
                logging.info("Found field_validations key later in response, attempting repair...")
                # Try to extract just the field_validations part
                start_idx = response_text.find('"field_validations"')
                if start_idx > 0:
                    response_text = '{"' + response_text[start_idx:]
            else:
                return None
        
        # Enhanced parsing to handle multiple bracket patterns
        import re
        
        lines = response_text.split('\n')
        field_validations = []
        current_object_lines = []
        brace_count = 0
        inside_validation = False
        in_field_validations_array = False
        
        for line_num, line in enumerate(lines):
            # Check if we're entering the field_validations array
            if '"field_validations":' in line:
                in_field_validations_array = True
                continue
                
            if not in_field_validations_array:
                continue
                
            # Look for object start - either opening brace alone or with field_id/field_name
            line_stripped = line.strip()
            if line_stripped == '{' or ('{' in line and ('"field_id"' in line or '"field_name"' in line)):
                # Start of a new field validation object
                if line_stripped == '{':
                    current_object_lines = [line]
                else:
                    current_object_lines = [line]
                inside_validation = True
                brace_count = line.count('{') - line.count('}')
            elif inside_validation:
                current_object_lines.append(line)
                brace_count += line.count('{') - line.count('}')
                
                # Check if we've completed this object (brace count is 0 and line ends with })
                if brace_count == 0 and (line.strip().endswith('}') or line.strip().endswith('},')):
                    # We have a complete field validation object
                    complete_object = '\n'.join(current_object_lines)
                    try:
                        # Try to parse this individual object to ensure it's valid
                        obj_json = complete_object.strip()
                        if obj_json.endswith(','):
                            obj_json = obj_json[:-1]  # Remove trailing comma
                        
                        # Wrap in a test structure to validate
                        test_json = '{"test": ' + obj_json + '}'
                        parsed_test = json.loads(test_json)
                        
                        # Validate it has required fields
                        test_obj = parsed_test['test']
                        if 'field_id' in test_obj or 'field_name' in test_obj:
                            field_validations.append(complete_object.strip())
                            logging.info(f"Found complete field validation object #{len(field_validations)}")
                        else:
                            logging.warning("Field validation object missing required keys")
                            
                    except json.JSONDecodeError as e:
                        logging.warning(f"Skipping invalid field validation object: {str(e)[:100]}...")
                    
                    # Reset for next object
                    current_object_lines = []
                    inside_validation = False
                    brace_count = 0
        
        if field_validations:
            # Build a proper JSON structure with all complete field validations
            repaired = '{\n  "field_validations": [\n'
            
            for i, validation in enumerate(field_validations):
                # Clean up the validation object
                clean_validation = validation.strip()
                if clean_validation.endswith(','):
                    clean_validation = clean_validation[:-1]
                
                # Add proper indentation
                indented_validation = '\n'.join('    ' + line for line in clean_validation.split('\n'))
                repaired += indented_validation
                
                # Add comma if not the last item
                if i < len(field_validations) - 1:
                    repaired += ','
                repaired += '\n'
            
            repaired += '  ]\n}'
            
            # Test if the repair worked
            final_parsed = json.loads(repaired)
            logging.info(f"✅ Successfully repaired truncated JSON - preserved {len(field_validations)} complete field validations")
            
            # Additional validation
            if 'field_validations' in final_parsed and isinstance(final_parsed['field_validations'], list):
                logging.info(f"✅ Repair validation passed - {len(final_parsed['field_validations'])} field validations ready")
                return repaired
            else:
                logging.error("❌ Repair validation failed - invalid structure")
                return None
        else:
            logging.warning("No complete field validation objects found in truncated response")
            return None
        
    except Exception as e:
        logging.error(f"JSON repair failed: {e}")
        import traceback
        logging.error(f"Repair traceback: {traceback.format_exc()}")
        return None

def step1_extract_from_documents(
    documents: List[Dict[str, Any]], 
    project_schema: Dict[str, Any],
    extraction_rules: Optional[List[Dict[str, Any]]] = None,
    knowledge_documents: Optional[List[Dict[str, Any]]] = None,
    session_name: str = "contract",
    validated_data_context: Optional[Dict[str, Any]] = None,
    extraction_notes: str = "",
    is_subsequent_upload: bool = False
) -> ExtractionResult:
    """
    STEP 1: Extract data from documents using AI
    
    Args:
        documents: List of document objects with file_content, file_name, mime_type
        project_schema: Schema definition with schema_fields and collections
        session_name: Name for the main object (default: "contract")
    
    Returns:
        ExtractionResult with extracted JSON data
    """
    try:
        # Check for API key
        api_key = os.environ.get("GEMINI_API_KEY")
        if not api_key:
            return ExtractionResult(success=False, error_message="GEMINI_API_KEY not found")
        
        # Import Google AI modules
        import google.generativeai as genai
        genai.configure(api_key=api_key)
        
        # Configure with no timeout constraints for large responses
        model = genai.GenerativeModel('gemini-2.0-flash-exp')
        
        logging.info(f"STEP 1: Starting extraction for {len(documents)} documents")
        
        # Debug document content
        for i, doc in enumerate(documents):
            content = doc.get('file_content', '')
            logging.info(f"Document {i}: {doc.get('file_name', 'Unknown')} - Content length: {len(content)}")
            if 'Active Deferred' in content or 'Code Meanings' in content:
                logging.info(f"Document {i} contains code meanings - content preview: {content[:300]}...")
            elif content:
                logging.info(f"Document {i} preview: {content[:200]}...")
        
        # First, identify global extraction rules that apply to all fields
        global_rules = []
        field_specific_rules = {}
        if extraction_rules:
            for rule in extraction_rules:
                rule_target = rule.get('targetField', '')
                rule_content = rule.get('ruleContent', '')
                
                # Check if this is a global rule (empty targetField or 'All Fields')
                if not rule_target or rule_target == '' or rule_target == 'All Fields':
                    global_rules.append(rule_content)
                else:
                    # Handle multiple target fields (comma-separated)
                    if ',' in rule_target:
                        targets = [t.strip() for t in rule_target.split(',')]
                        for target in targets:
                            if target not in field_specific_rules:
                                field_specific_rules[target] = []
                            field_specific_rules[target].append(rule_content)
                    else:
                        if rule_target not in field_specific_rules:
                            field_specific_rules[rule_target] = []
                        field_specific_rules[rule_target].append(rule_content)
        
        # Build verified context section
        verified_context_text = ""
        verified_field_ids = set()
        highest_collection_indices = {}
        
        # Use collection_record_counts from server if available
        collection_record_counts = input_data.get('collection_record_counts', {})
        logging.info(f"INPUT: collection_record_counts: {collection_record_counts}")
        logging.info(f"INPUT: validated_data_context contains {len(validated_data_context) if validated_data_context else 0} items")
        if validated_data_context:
            logging.info(f"SAMPLE verified context items: {list(validated_data_context.keys())[:5]}")
        
        # Initialize collections tracking
        verified_fields = []
        verified_collections = {}
        existing_collections = {}  # Track all existing collection items for field targeting
        
        if validated_data_context:
            
            # Extract verified fields and collections from validated_data_context
            for field_key, field_data in validated_data_context.items():
                field_id = field_data.get('field_id')
                validation_status = field_data.get('validation_status', 'unverified')
                field_name = field_data.get('field_name', 'Unknown')
                extracted_value = field_data.get('extracted_value', '')
                
                # Track all existing collection items (both verified and unverified) for field targeting
                if field_data.get('validation_type') == 'collection_property':
                    collection_name = field_data.get('collection_name', '')
                    record_index = field_data.get('record_index', 0)
                    
                    if collection_name not in existing_collections:
                        existing_collections[collection_name] = {}
                    if record_index not in existing_collections[collection_name]:
                        existing_collections[collection_name][record_index] = []
                    
                    existing_collections[collection_name][record_index].append({
                        'field_id': field_id,
                        'field_name': field_name,
                        'extracted_value': extracted_value,
                        'validation_status': validation_status
                    })
                    
                    # Track highest index for each collection
                    if collection_name not in highest_collection_indices:
                        highest_collection_indices[collection_name] = record_index
                    else:
                        highest_collection_indices[collection_name] = max(highest_collection_indices[collection_name], record_index)
                
                # Only treat as verified if status is actually verified
                if validation_status == 'verified':
                    verified_field_ids.add(field_id)
                    
                    if field_data.get('validation_type') == 'schema_field':
                        verified_fields.append(f"- **{field_name}**: {extracted_value}")
                    elif field_data.get('validation_type') == 'collection_property':
                        collection_name = field_data.get('collection_name', '')
                        record_index = field_data.get('record_index', 0)
                        
                        if collection_name not in verified_collections:
                            verified_collections[collection_name] = {}
                        if record_index not in verified_collections[collection_name]:
                            verified_collections[collection_name][record_index] = []
                        
                        verified_collections[collection_name][record_index].append(f"  - {field_name}: {extracted_value}")
            
            # Log existing collection structure for debugging
            for collection_name, records in existing_collections.items():
                record_count = len(records)
                logging.info(f"EXISTING COLLECTION: {collection_name} has {record_count} existing records (indices: {sorted(records.keys())})")
                
                # Log some field details for debugging
                for record_index in sorted(list(records.keys())[:3]):  # Show first 3 records
                    fields = records[record_index]
                    field_info = [f"{f['field_name']}={f['extracted_value'] or 'EMPTY'}" for f in fields]
                    logging.info(f"  Record {record_index}: {', '.join(field_info[:3])}...")  # Show first 3 fields
                
                # Override with server-provided collection record counts if available
                for collection_name, count in collection_record_counts.items():
                    if count > 0:
                        highest_collection_indices[collection_name] = count - 1  # Convert count to highest index
                        logging.info(f"OVERRIDE: Using server count for {collection_name}: {count} records (highest index: {count - 1})")
            
            # Build context text
            if verified_fields or verified_collections:
                verified_context_text = "**VERIFIED FIELDS (DO NOT RE-EXTRACT OR MODIFY):**\n"
                
                if verified_fields:
                    verified_context_text += "\n**Schema Fields (LOCKED):**\n" + "\n".join(verified_fields) + "\n"
                
                if verified_collections:
                    verified_context_text += "\n**Collection Items (LOCKED):**\n"
                    for collection_name, items in verified_collections.items():
                        highest_index = highest_collection_indices.get(collection_name, 0)
                        next_index = highest_index + 1
                        verified_context_text += f"\n**{collection_name}** (Verified Items 0-{highest_index}):\n"
                        for record_index in sorted(items.keys()):
                            verified_context_text += f"Item {record_index}:\n" + "\n".join(items[record_index]) + "\n"
                        verified_context_text += f"**NEXT NEW ITEMS START AT INDEX {next_index}**\n"
                        logging.info(f"CONTEXT: {collection_name} highest verified index: {highest_index}, next index: {next_index}")
                
                verified_context_text += "\n**CRITICAL**: Do not include any of the above verified fields in your extraction response. Use them only as context for understanding the session.\n"
            else:
                verified_context_text = "No verified fields yet - this is the initial extraction."
                logging.info("CONTEXT: No verified fields found - initial extraction")

        # Build schema fields section for the imported prompt
        schema_fields_text = ""
        
        # Add schema fields with descriptions for AI guidance  
        if project_schema.get("schema_fields"):
            for field in project_schema["schema_fields"]:
                field_name = field['fieldName']
                field_type = field['fieldType']
                field_description = field.get('description', '')
                field_id = field['id']
                
                # Skip verified fields
                if field_id in verified_field_ids:
                    continue
                
                # Collect all applicable extraction rules for this field
                applicable_rules = []
                
                # Add global rules first (these apply to ALL fields)
                applicable_rules.extend([f"RULE: {rule}" for rule in global_rules])
                
                # Add field-specific rules
                if field_name in field_specific_rules:
                    applicable_rules.extend([f"RULE: {rule}" for rule in field_specific_rules[field_name]])
                
                # Combine description with rules
                full_instruction = field_description or 'Extract this field from the documents'
                if applicable_rules:
                    full_instruction += " | " + " | ".join(applicable_rules)
                
                # Include field ID in the prompt for AI reference
                schema_fields_text += f"\n- **{field_name}** (ID: {field_id}, {field_type}): {full_instruction}"
        
        # Build collections section for the imported prompt
        collections_text = ""
        if project_schema.get("collections"):
            for collection in project_schema["collections"]:
                collection_name = collection.get('collectionName', collection.get('objectName', ''))
                collection_description = collection.get('description', '')
                
                # Check if this collection has verified items
                verified_count = highest_collection_indices.get(collection_name, -1) + 1 if collection_name in highest_collection_indices else 0
                next_index = verified_count
                
                # Check if this collection has existing items for field targeting
                existing_count = 0
                if collection_name in existing_collections:
                    existing_count = len(existing_collections[collection_name])
                    logging.info(f"FIELD TARGETING: {collection_name} has {existing_count} existing collection items")
                
                # Find applicable extraction rules for this collection
                applicable_rules = []
                if extraction_rules:
                    for rule in extraction_rules:
                        rule_target = rule.get('targetField', [])
                        if isinstance(rule_target, list):
                            if collection_name in rule_target or 'All Fields' in rule_target:
                                applicable_rules.append(f"RULE: {rule.get('ruleContent', '')}")
                        elif collection_name == rule_target or rule_target == 'All Fields':
                            applicable_rules.append(f"RULE: {rule.get('ruleContent', '')}")
                
                full_instruction = collection_description or 'Extract array of these objects'
                if applicable_rules:
                    full_instruction += " | " + " | ".join(applicable_rules)
                
                collections_text += f"\n- **{collection_name}**: {full_instruction}"
                
                # Add field targeting context for existing collection items
                if existing_count > 0:
                    collections_text += f"\n  **FIELD TARGETING MODE**: This collection has {existing_count} existing items (record indices 0-{existing_count-1})"
                    collections_text += f"\n  **UPDATE EXISTING ITEMS**: Fill in missing field values for existing collection items. Do NOT create new items unless extracting additional data not present in existing records."
                    collections_text += f"\n  **MAINTAIN RECORD INDICES**: When updating existing items, keep the same record_index values (0, 1, 2, etc.) as the existing collection structure."
                
                # Add verified collection context
                elif verified_count > 0:
                    collections_text += f"\n  **VERIFIED ITEMS EXIST**: Items 0-{verified_count-1} are already verified and locked. Do NOT re-extract these."
                    collections_text += f"\n  **START NEW ITEMS AT INDEX {next_index}**: Any new items must use record_index {next_index}, {next_index+1}, {next_index+2}, etc."
                
                # Add explicit instructions for list/collection items
                collections_text += f"\n  **CRITICAL FOR {collection_name}**: Find ALL instances in the documents. Create one collection item per unique instance found. Each item should have a separate record_index ({next_index}, {next_index+1}, {next_index+2}, etc.)."
                collections_text += f"\n  **TABLE EXTRACTION**: If {collection_name} items appear in a table, extract EVERY ROW from that table, not just 2-3 examples. Count all rows and extract all data."
                collections_text += f"\n  **NUMBERED SECTIONS**: If {collection_name} matches a document section name, find ALL numbered subsections (e.g., 2.3.1, 2.3.2, 2.3.3, etc.) and extract each one as a separate collection item."
                collections_text += f"\n  **MARKDOWN TABLES**: Recognize markdown table format with | separators. Extract ALL data rows (excluding headers) as separate collection items - if table has 10 rows, extract all 10."
                collections_text += f"\n  **COUNT ALL ITEMS**: If you see numbered items 2.3.1, 2.3.2, 2.3.3... keep going until you reach the end (e.g., 2.3.10) and extract EVERY SINGLE ONE as separate collection items."
                collections_text += f"\n  **NO TRUNCATION**: Include ALL found items in your JSON response - do not truncate or limit the output even if there are many items."
                
                properties = collection.get("properties", [])
                if properties:
                    collections_text += f"\n  Properties for each {collection_name} item:"
                    for prop in properties:
                        prop_name = prop.get('propertyName', '')
                        prop_type = prop.get('propertyType', 'TEXT')
                        prop_description = prop.get('description', '')
                        
                        # Find applicable extraction rules for this property
                        prop_rules = []
                        
                        # Add global rules first (these apply to ALL fields)
                        prop_rules.extend([f"RULE: {rule}" for rule in global_rules])
                        
                        # Add property-specific rules
                        arrow_notation = f"{collection_name} --> {prop_name}"
                        full_prop_name = f"{collection_name}.{prop_name}"
                        
                        # Check field_specific_rules for various naming patterns
                        for pattern in [arrow_notation, full_prop_name, prop_name]:
                            if pattern in field_specific_rules:
                                prop_rules.extend([f"RULE: {rule}" for rule in field_specific_rules[pattern]])
                        
                        prop_instruction = prop_description or 'Extract this property'
                        if prop_rules:
                            prop_instruction += " | " + " | ".join(prop_rules)
                            logging.info(f"RULE MATCH: {collection_name} --> {prop_name} matched rules: {[rule.get('ruleName') for rule in extraction_rules if arrow_notation in str(rule.get('targetField', []))]}")
                        
                        # Add choice options for CHOICE fields
                        if prop_type == 'CHOICE' and prop.get('choiceOptions'):
                            choice_text = f"The output should be one of the following choices: {'; '.join(prop['choiceOptions'])}."
                            prop_instruction = prop_instruction + " | " + choice_text if prop_instruction else choice_text
                            
                        # Include property ID in the prompt for AI reference
                        prop_id = prop['id']
                        collections_text += f"\n  * **{prop_name}** (ID: {prop_id}, {prop_type}): {prop_instruction}"
        
        # Generate JSON schema section showing exact field mappings
        json_schema_section = ""
        
        # First, create a dedicated Knowledge Documents section
        if knowledge_documents:
            json_schema_section += "\n## KNOWLEDGE DOCUMENTS:\n"
            json_schema_section += "Use the following knowledge documents for context and validation:\n\n"
            for doc in knowledge_documents:
                display_name = doc.get('displayName', doc.get('fileName', 'Unknown Document'))
                description = doc.get('description', '')
                content = doc.get('content', '')
                target_field = doc.get('targetField', '')
                
                json_schema_section += f"### {display_name}\n"
                if description:
                    json_schema_section += f"**Description**: {description}\n"
                if target_field:
                    json_schema_section += f"**Applies to**: {target_field}\n"
                else:
                    json_schema_section += f"**Applies to**: All fields\n"
                json_schema_section += f"**Content**:\n```\n{content}\n```\n\n"
        
        # Add dedicated Extraction Rules section
        if extraction_rules:
            json_schema_section += "\n## EXTRACTION RULES:\n"
            json_schema_section += "The following extraction rules modify field behavior:\n\n"
            for i, rule in enumerate(extraction_rules):
                rule_name = rule.get('ruleName', f'Rule {i+1}')
                rule_content = rule.get('ruleContent', '')
                rule_target = rule.get('targetField', [])
                
                json_schema_section += f"### {rule_name}\n"
                if rule_content:
                    json_schema_section += f"**Rule**: {rule_content}\n"
                if rule_target:
                    if isinstance(rule_target, list):
                        json_schema_section += f"**Applies to**: {', '.join(rule_target)}\n"
                    else:
                        json_schema_section += f"**Applies to**: {rule_target}\n"
                else:
                    json_schema_section += f"**Applies to**: All fields\n"
                json_schema_section += "\n"
        
        # Add schema fields JSON format
        if project_schema.get("schema_fields"):
            json_schema_section += "\n## SCHEMA FIELDS JSON FORMAT:\n"
            json_schema_section += "```json\n{\n  \"schema_fields\": [\n"
            for i, field in enumerate(project_schema["schema_fields"]):
                field_id = field['id']
                field_name = field['fieldName']
                field_type = field['fieldType']
                field_description = field.get('description', '')
                auto_verify_threshold = field.get('autoVerificationConfidence', 80)
                
                # Find applicable extraction rules for this field
                applicable_rules = []
                
                # Add global rules first (these apply to ALL fields)
                applicable_rules.extend(global_rules)
                
                # Add field-specific rules
                if field_name in field_specific_rules:
                    applicable_rules.extend(field_specific_rules[field_name])

                # Find applicable knowledge documents for this field (just names for reference)
                applicable_knowledge_names = []
                if knowledge_documents:
                    for doc in knowledge_documents:
                        doc_target = doc.get('targetField', '')
                        display_name = doc.get('displayName', doc.get('fileName', 'Unknown Document'))
                        # If no target field specified, apply to all fields
                        if not doc_target or doc_target == '' or doc_target is None:
                            applicable_knowledge_names.append(display_name)
                        elif isinstance(doc_target, list):
                            if field_name in doc_target or 'All Fields' in doc_target:
                                applicable_knowledge_names.append(display_name)
                        elif field_name == doc_target or doc_target == 'All Fields':
                            applicable_knowledge_names.append(display_name)
                
                # Combine description with rules
                full_instruction = field_description or 'Extract this field from the documents'
                if applicable_rules:
                    full_instruction += " | RULE: " + " | RULE: ".join(applicable_rules)
                
                json_schema_section += f"    {{\n"
                json_schema_section += f"      \"field_id\": \"{field_id}\",\n"
                json_schema_section += f"      \"field_name\": \"{field_name}\",\n"
                json_schema_section += f"      \"field_type\": \"{field_type}\",\n"
                json_schema_section += f"      \"description\": \"{full_instruction}\",\n"
                json_schema_section += f"      \"auto_verification_confidence\": {auto_verify_threshold}"
                if field_type == 'CHOICE' and field.get('choiceOptions'):
                    json_schema_section += f",\n      \"choices\": {field['choiceOptions']}"
                
                # Add extraction rules section
                if applicable_rules:
                    rules_list = [rule.replace('"', '\\"') for rule in applicable_rules]
                    json_schema_section += f",\n      \"extraction_rules\": {rules_list}"
                else:
                    json_schema_section += f",\n      \"extraction_rules\": []"
                
                # Add knowledge documents section (just document names)
                if applicable_knowledge_names:
                    json_schema_section += f",\n      \"knowledge_documents\": {json.dumps(applicable_knowledge_names)}"
                else:
                    json_schema_section += f",\n      \"knowledge_documents\": []"
                
                json_schema_section += "\n    }"
                if i < len(project_schema["schema_fields"]) - 1:
                    json_schema_section += ","
                json_schema_section += "\n"
            json_schema_section += "  ]\n}\n```\n"
        
        # Add collections JSON format
        if project_schema.get("collections"):
            json_schema_section += "\n## COLLECTIONS JSON FORMAT:\n"
            json_schema_section += "```json\n{\n  \"collections\": [\n"
            for i, collection in enumerate(project_schema["collections"]):
                collection_name = collection.get('collectionName', collection.get('objectName', ''))
                collection_description = collection.get('description', '')
                
                # Find applicable extraction rules for this collection
                applicable_rules = []
                
                # Add global rules first (these apply to ALL fields)
                applicable_rules.extend(global_rules)
                
                # Add collection-specific rules
                if collection_name in field_specific_rules:
                    applicable_rules.extend(field_specific_rules[collection_name])

                # Find applicable knowledge documents for this collection (just names for reference)
                applicable_coll_knowledge_names = []
                if knowledge_documents:
                    for doc in knowledge_documents:
                        doc_target = doc.get('targetField', '')
                        display_name = doc.get('displayName', doc.get('fileName', 'Unknown Document'))
                        # If no target field specified, apply to all fields/collections
                        if not doc_target or doc_target == '' or doc_target is None:
                            applicable_coll_knowledge_names.append(display_name)
                        elif isinstance(doc_target, list):
                            if collection_name in doc_target or 'All Fields' in doc_target:
                                applicable_coll_knowledge_names.append(display_name)
                        elif collection_name == doc_target or doc_target == 'All Fields':
                            applicable_coll_knowledge_names.append(display_name)
                
                full_instruction = collection_description or 'Extract array of these objects'
                if applicable_rules:
                    full_instruction += " | RULE: " + " | RULE: ".join(applicable_rules)
                
                json_schema_section += f"    {{\n"
                json_schema_section += f"      \"collection_name\": \"{collection_name}\",\n"
                json_schema_section += f"      \"description\": \"{full_instruction}\",\n"
                
                # Add extraction rules section for collections
                if applicable_rules:
                    coll_rules_list = [rule.replace('"', '\\"') for rule in applicable_rules]
                    json_schema_section += f"      \"extraction_rules\": {coll_rules_list},\n"
                else:
                    json_schema_section += f"      \"extraction_rules\": [],\n"
                
                # Add knowledge documents section for collections (just document names)
                if applicable_coll_knowledge_names:
                    json_schema_section += f"      \"knowledge_documents\": {json.dumps(applicable_coll_knowledge_names)},\n"
                else:
                    json_schema_section += f"      \"knowledge_documents\": [],\n"
                
                json_schema_section += f"      \"properties\": [\n"
                
                properties = collection.get("properties", [])
                for j, prop in enumerate(properties):
                    prop_name = prop.get('propertyName', '')
                    prop_type = prop.get('propertyType', 'TEXT')
                    prop_description = prop.get('description', '')
                    prop_id = prop['id']
                    prop_auto_verify_threshold = prop.get('autoVerificationConfidence', 80)
                    
                    # Find applicable extraction rules for this property
                    prop_rules = []
                    
                    # Add global rules first (these apply to ALL fields)
                    prop_rules.extend(global_rules)
                    
                    # Add property-specific rules
                    arrow_notation = f"{collection_name} --> {prop_name}"
                    full_prop_name = f"{collection_name}.{prop_name}"
                    
                    # Check field_specific_rules for various naming patterns
                    for pattern in [arrow_notation, full_prop_name, prop_name]:
                        if pattern in field_specific_rules:
                            prop_rules.extend(field_specific_rules[pattern])

                    # Find applicable knowledge documents for this property (just names for reference)
                    prop_knowledge_names = []
                    if knowledge_documents:
                        for doc in knowledge_documents:
                            doc_target = doc.get('targetField', '')
                            arrow_notation = f"{collection_name} --> {prop_name}"
                            full_prop_name = f"{collection_name}.{prop_name}"
                            display_name = doc.get('displayName', doc.get('fileName', 'Unknown Document'))
                            
                            # If no target field specified, apply to all fields/properties
                            if not doc_target or doc_target == '' or doc_target is None:
                                prop_knowledge_names.append(display_name)
                            elif isinstance(doc_target, list):
                                if (arrow_notation in doc_target or 
                                    full_prop_name in doc_target or 
                                    prop_name in doc_target or 
                                    'All Fields' in doc_target):
                                    prop_knowledge_names.append(display_name)
                            elif (arrow_notation == doc_target or 
                                  full_prop_name == doc_target or 
                                  prop_name == doc_target or 
                                  doc_target == 'All Fields'):
                                prop_knowledge_names.append(display_name)
                    
                    prop_instruction = prop_description or 'Extract this property'
                    if prop_rules:
                        prop_instruction += " | RULE: " + " | RULE: ".join(prop_rules)
                    
                    json_schema_section += f"        {{\n"
                    json_schema_section += f"          \"property_id\": \"{prop_id}\",\n"
                    json_schema_section += f"          \"property_name\": \"{prop_name}\",\n"
                    json_schema_section += f"          \"property_type\": \"{prop_type}\",\n"
                    json_schema_section += f"          \"description\": \"{prop_instruction}\",\n"
                    json_schema_section += f"          \"auto_verification_confidence\": {prop_auto_verify_threshold}"
                    if prop_type == 'CHOICE' and prop.get('choiceOptions'):
                        json_schema_section += f",\n          \"choices\": {prop['choiceOptions']}"
                    
                    # Add extraction rules section for properties
                    if prop_rules:
                        prop_rules_list = [rule.replace('"', '\\"') for rule in prop_rules]
                        json_schema_section += f",\n          \"extraction_rules\": {prop_rules_list}"
                    else:
                        json_schema_section += f",\n          \"extraction_rules\": []"
                    
                    # Add knowledge documents section for properties (just document names)
                    if prop_knowledge_names:
                        json_schema_section += f",\n          \"knowledge_documents\": {json.dumps(prop_knowledge_names)}"
                    else:
                        json_schema_section += f",\n          \"knowledge_documents\": []"
                    
                    json_schema_section += "\n        }"
                    if j < len(properties) - 1:
                        json_schema_section += ","
                    json_schema_section += "\n"
                
                json_schema_section += "      ]\n    }"
                if i < len(project_schema["collections"]) - 1:
                    json_schema_section += ","
                json_schema_section += "\n"
            json_schema_section += "  ]\n}\n```\n"
        
        # Get additional instructions from input data
        additional_instructions = input_data.get('additional_instructions', '')
        if not additional_instructions:
            additional_instructions = "No additional instructions provided."
        
        # Handle existing collection records exclusion
        existing_records_text = ""
        existing_collection_records = input_data.get('existing_collection_records', {})
        if existing_collection_records:
            existing_records_text = "\n\n**IMPORTANT - EXISTING COLLECTION RECORDS TO SKIP:**\n"
            for collection_name, record_indexes in existing_collection_records.items():
                if record_indexes:
                    existing_records_text += f"- Collection '{collection_name}': Skip record indexes {record_indexes} (already extracted)\n"
            existing_records_text += "\nYou MUST extract NEW records only. Do not extract the same record indexes listed above. Start from the next available record index.\n"
            logging.info(f"EXCLUSION LOGIC: Skipping existing records: {existing_collection_records}")
        
        # Use the imported prompt template with our schema, collections, verified context, and JSON schema
        full_prompt = json_schema_section + "\n" + EXTRACTION_PROMPT.format(
            verified_context=verified_context_text,
            schema_fields=schema_fields_text,
            collections=collections_text,
            additional_instructions=additional_instructions + existing_records_text
        )
        
        # Generate field validation JSON structure
        def generate_field_validation_example():
            json_lines = ['{"field_validations": [']
            
            # Add schema fields with proper field validation structure
            if project_schema.get("schema_fields"):
                for i, field in enumerate(project_schema["schema_fields"]):
                    field_id = field['id']
                    field_name = field['fieldName']
                    field_type = field['fieldType']
                    field_description = field.get('description', '')
                    
                    # Determine example value based on field type
                    if field_type == 'NUMBER':
                        if 'parties' in field_name.lower():
                            example_value = '33'
                        elif 'nda' in field_name.lower():
                            example_value = '8'
                        else:
                            example_value = '42'
                    elif field_type == 'DATE':
                        example_value = '2024-01-15'
                    elif field_type == 'CHOICE' and field.get('choiceOptions'):
                        example_value = field["choiceOptions"][0]
                    else:  # TEXT
                        example_value = 'Extracted Text Value'
                    
                    # Build AI reasoning
                    reasoning = f"Extracted from document analysis"
                    if field_description:
                        reasoning += f" - {field_description}"
                    
                    json_lines.append('  {')
                    json_lines.append(f'    "field_id": "{field_id}",')
                    json_lines.append(f'    "validation_type": "schema_field",')
                    json_lines.append(f'    "data_type": "{field_type}",')
                    json_lines.append(f'    "field_name": "{field_name}",')
                    json_lines.append(f'    "extracted_value": "{example_value}",')
                    json_lines.append(f'    "confidence_score": 0.95,')
                    json_lines.append(f'    "validation_status": "unverified",')
                    json_lines.append(f'    "ai_reasoning": "Found in document section X - {field_description}"')
                    json_lines.append('  }' + (',' if i < len(project_schema["schema_fields"]) - 1 or project_schema.get("collections") else ''))
            
            # Add collection properties with proper field validation structure
            if project_schema.get("collections"):
                # Track items to determine which is the last one for proper JSON formatting
                all_collection_items = []
                
                for collection in project_schema["collections"]:
                    collection_name = collection.get('collectionName', collection.get('objectName', ''))
                    properties = collection.get("properties", [])
                    
                    # Only process collections that have properties
                    if not properties:
                        continue
                        
                    # Show minimal examples for all collections within 100-record limit
                    # Calculate available space for collection examples
                    schema_field_count = len(project_schema.get("schema_fields", []))
                    remaining_space = 100 - schema_field_count
                    total_collections = len(project_schema.get("collections", []))
                    
                    # Distribute remaining space among collections, minimum 1 example each
                    if total_collections > 0:
                        properties_per_collection = sum(len(c.get('properties', [])) for c in project_schema.get("collections", []))
                        if properties_per_collection > 0:
                            max_examples_per_collection = max(1, remaining_space // (properties_per_collection))
                            example_count = min(2, max_examples_per_collection)  # Standard example count, capped by space
                        else:
                            example_count = 1
                    else:
                        example_count = 2
                    
                    # Start from next available index for this collection
                    start_index = highest_collection_indices.get(collection_name, -1) + 1
                    
                    # Generate all collection items for this collection
                    for record_index in range(start_index, start_index + example_count):
                        for prop_index, prop in enumerate(properties):
                            prop_id = prop['id']
                            prop_name = prop['propertyName']
                            prop_type = prop['propertyType']
                            
                            # Determine example value based on field type
                            if prop_type == 'CHOICE' and prop.get('choiceOptions'):
                                example_value = prop["choiceOptions"][0]
                            elif prop_type == 'NUMBER':
                                example_value = '100'
                            elif prop_type == 'DATE':
                                example_value = '2024-01-15'
                            else:  # TEXT
                                example_value = 'Extracted Value'
                            
                            field_name_with_index = f"{collection_name}.{prop_name}[{record_index}]"
                            
                            # Store item info for later JSON generation
                            all_collection_items.append({
                                'prop_id': prop_id,
                                'prop_type': prop_type,
                                'field_name_with_index': field_name_with_index,
                                'collection_name': collection_name,
                                'example_value': example_value,
                                'record_index': record_index,
                                'prop_name': prop_name
                            })
                
                # Generate JSON for all collection items
                for i, item in enumerate(all_collection_items):
                    json_lines.append('  {')
                    json_lines.append(f'    "field_id": "{item["prop_id"]}",')
                    json_lines.append(f'    "validation_type": "collection_property",')
                    json_lines.append(f'    "data_type": "{item["prop_type"]}",')
                    json_lines.append(f'    "field_name": "{item["field_name_with_index"]}",')
                    json_lines.append(f'    "collection_name": "{item["collection_name"]}",')
                    json_lines.append(f'    "extracted_value": "{item["example_value"]}",')
                    json_lines.append(f'    "confidence_score": 0.95,')
                    json_lines.append(f'    "validation_status": "unverified",')
                    json_lines.append(f'    "ai_reasoning": "Found {item["collection_name"]} item {item["record_index"] + 1} with {item["prop_name"]} value in document",')
                    json_lines.append(f'    "record_index": {item["record_index"]}')
                    
                    # Check if this is the last item
                    is_last = (i == len(all_collection_items) - 1)
                    json_lines.append('  }' + ('' if is_last else ','))
            
            json_lines.append(']}')
            return '\n'.join(json_lines)
        
        dynamic_example = generate_field_validation_example()
        logging.info(f"Generated field validation example with {len(extraction_rules or [])} extraction rules")
        logging.info(f"Dynamic example preview (first 500 chars): {dynamic_example[:500]}...")
        
        # Handle validated data context for subsequent uploads
        validated_context = ""
        if is_subsequent_upload and validated_data_context:
            logging.info(f"Processing validated data context with {len(validated_data_context)} existing fields (subsequent upload mode)")
            validated_context = "\n\n## EXISTING VALIDATED DATA (READ-ONLY CONTEXT):\n"
            validated_context += "The following data has been previously validated and is provided for context. DO NOT include these fields in your output:\n\n"
            
            # Group by collection for better organization
            collections_data = {}
            schema_data = {}
            
            for field_name, context in validated_data_context.items():
                if context.get('collection'):
                    collection_name = context['collection']
                    if collection_name not in collections_data:
                        collections_data[collection_name] = {}
                    
                    record_index = context.get('recordIndex', 0)
                    if record_index not in collections_data[collection_name]:
                        collections_data[collection_name][record_index] = {}
                    
                    # Extract property name from field_name like "Codes.Code Name[0]"
                    property_name = field_name.split('.')[1].split('[')[0] if '.' in field_name else field_name
                    collections_data[collection_name][record_index][property_name] = context
                else:
                    schema_data[field_name] = context
            
            # Display schema fields
            if schema_data:
                validated_context += "**Schema Fields (Validated):**\n"
                for field_name, context in schema_data.items():
                    validated_context += f"- {field_name}: \"{context['value']}\" (Confidence: {context.get('confidence', 0)}%)\n"
                    if context.get('reasoning'):
                        validated_context += f"  Reasoning: {context['reasoning']}\n"
                validated_context += "\n"
            
            # Display collection data
            for collection_name, records in collections_data.items():
                validated_context += f"**{collection_name} Collection (Validated Records):**\n"
                for record_index, properties in records.items():
                    validated_context += f"Record {record_index}:\n"
                    for prop_name, context in properties.items():
                        validated_context += f"  - {prop_name}: \"{context['value']}\"\n"
                        if context.get('reasoning'):
                            validated_context += f"    Reasoning: {context['reasoning']}\n"
                validated_context += "\n"
            
            validated_context += "## EXTRACTION INSTRUCTIONS FOR SUBSEQUENT UPLOAD:\n"
            validated_context += "- **FOCUS ON MISSING DATA**: Only extract data for fields that are NOT shown above\n"
            validated_context += "- **MAINTAIN COLLECTION INTEGRITY**: For collections, ensure new records use the next available index\n"
            validated_context += "- **EXTRACT FROM NEW DOCUMENT**: Use the uploaded document content to fill missing information\n"
            validated_context += "- **COMPREHENSIVE EXTRACTION**: Extract detailed explanations, not abbreviated summaries\n\n"

        # The imported prompt already contains all the necessary instructions
        # Just add document verification and choice field handling specific to this run
        # Add extraction notes if provided
        extraction_notes_context = ""
        if extraction_notes:
            extraction_notes_context = f"\n\n## EXTRACTION NOTES:\n{extraction_notes}\n"

        full_prompt += f"""
{validated_context}{extraction_notes_context}
DOCUMENT VERIFICATION: Confirm you processed all {len(documents)} documents: {[doc.get('file_name', 'Unknown') for doc in documents]}

**RESPONSE OPTIMIZATION**: Aim for comprehensive extraction while maintaining quality:
- Extract ALL relevant data found in the documents
- For large datasets, prioritize complete records over partial ones
- Ensure all collection items are fully populated with their properties
- Focus on accuracy and completeness rather than arbitrary limits

CHOICE FIELD HANDLING:
- For CHOICE fields, extract values from the specified choice options only
- If the document contains values not in the choice options, return null (do not block processing)
- Choice options are specified as "The output should be one of the following choices: ..."
- Example: For Yes/No choice, only return "Yes" or "No", never "true", "false", "1", "0", etc.

**CRITICAL FIELD ID REQUIREMENT**: Use the EXACT field_id values provided in the schema above (the UUID values after "ID: "). For example, if you see "Product/Service Specifications Met (ID: c3056038-5b32-4335-8772-a95c9bef307a, CHOICE)", use "c3056038-5b32-4335-8772-a95c9bef307a" as the field_id. Do not use field names, camelCase versions, or generate your own IDs.

**CRITICAL COLLECTION NAME REQUIREMENT**: For collection properties, you MUST include the "collection_name" field in each field validation object. Use the exact collection name from the schema (e.g., "Increase Rates").

REQUIRED OUTPUT FORMAT - Field Validation JSON Structure:
{dynamic_example}"""
        
        # STEP 1: ENHANCED DOCUMENT CONTENT EXTRACTION
        # Process documents in two phases: content extraction, then data extraction
        logging.info(f"=== STEP 1: DOCUMENT CONTENT EXTRACTION ===")
        logging.info(f"Processing {len(documents)} documents for content extraction")
        logging.info(f"Documents received: {[doc.get('file_name', 'Unknown') for doc in documents]}")
        logging.info(f"Documents data: {documents}")
        
        model = genai.GenerativeModel('gemini-2.5-pro')
        extracted_content_text = ""
        processed_docs = 0
        
        for doc in documents:
            file_content = doc['file_content']
            file_name = doc['file_name']
            mime_type = doc['mime_type']
            
            logging.info(f"STEP 1: Processing document: {file_name} ({mime_type})")
            
            # Handle document content - prioritize already extracted text content
            if isinstance(file_content, str) and not file_content.startswith('data:'):
                # Check only for explicit error messages (not short content length)
                if ("503 UNAVAILABLE" in file_content or 
                    "Error extracting text" in file_content or 
                    ("overloaded" in file_content.lower() and "PDF processing overloaded" in file_content)):
                    logging.warning(f"STEP 1: Detected failed extraction content for {file_name}, content: {file_content[:100]}...")
                    # Skip this document and log the issue
                    extracted_content_text += f"\n\n=== DOCUMENT: {file_name} ===\n[DOCUMENT CONTENT EXTRACTION FAILED - PDF processing overloaded. Please retry the extraction to process this {file_name} document properly.]"
                    processed_docs += 1
                    continue
                else:
                    # This is legitimate pre-extracted text content from the session
                    content_text = file_content
                    extracted_content_text += f"\n\n=== DOCUMENT: {file_name} ===\n{content_text}"
                    processed_docs += 1
                    logging.info(f"STEP 1: Using pre-extracted content from {file_name} ({len(content_text)} characters)")
                
            elif mime_type.startswith("text/"):
                # Handle plain text content
                if isinstance(file_content, str):
                    if file_content.startswith('data:'):
                        base64_content = file_content.split(',', 1)[1]
                        decoded_bytes = base64.b64decode(base64_content)
                        content_text = decoded_bytes.decode('utf-8', errors='ignore')
                    else:
                        content_text = file_content
                else:
                    content_text = file_content.decode('utf-8', errors='ignore')
                
                extracted_content_text += f"\n\n=== DOCUMENT: {file_name} ===\n{content_text}"
                processed_docs += 1
                logging.info(f"STEP 1: Extracted {len(content_text)} characters from text file {file_name}")
                
            else:
                # Use Gemini API for binary document content extraction (PDF, Word, Excel)
                try:
                    if isinstance(file_content, str) and file_content.startswith('data:'):
                        # Decode data URL
                        mime_part, base64_content = file_content.split(',', 1)
                        binary_content = base64.b64decode(base64_content)
                        
                        # Create specialized extraction prompts based on document type
                        if ('excel' in mime_type or 
                            'spreadsheet' in mime_type or 
                            'vnd.ms-excel' in mime_type or 
                            'vnd.openxmlformats-officedocument.spreadsheetml' in mime_type or
                            file_name.lower().endswith(('.xlsx', '.xls'))):
                            # Excel file - extract all sheets and tabular data
                            extraction_prompt = f"""Extract ALL content from this Excel file ({file_name}).

INSTRUCTIONS:
- Extract content from ALL worksheets/sheets in the workbook
- For each sheet, include the sheet name as a header
- Preserve table structure where possible using clear formatting
- Include all text, numbers, formulas results, and data
- If there are multiple sheets, clearly separate them with sheet names
- Format the output as readable structured text that preserves the original data organization

RETURN: Complete text content from all sheets in this Excel file."""

                        elif 'pdf' in mime_type or file_name.lower().endswith('.pdf'):
                            # PDF file - extract all text content with enhanced retry logic
                            extraction_prompt = f"""Extract ALL text content from this PDF document ({file_name}).

INSTRUCTIONS:
- Extract all readable text from every page
- Preserve document structure and formatting where possible
- Include headers, body text, tables, lists, and any other textual content
- Maintain logical flow and organization of information
- Focus on key data points and structured information

RETURN: Complete text content from this PDF document."""

                        elif ('word' in mime_type or 
                              'vnd.openxmlformats-officedocument.wordprocessingml' in mime_type or
                              'application/msword' in mime_type or
                              file_name.lower().endswith(('.docx', '.doc'))):
                            # Word document - extract all content
                            extraction_prompt = f"""Extract ALL content from this Word document ({file_name}).

INSTRUCTIONS:
- Extract all text content including body text, headers, footers, tables
- Preserve document structure and formatting where possible
- Include any embedded text, comments, or annotations
- Maintain logical organization of the content

RETURN: Complete text content from this Word document."""

                        else:
                            # Generic binary file extraction
                            extraction_prompt = f"""Extract all readable text content from this document ({file_name}).

INSTRUCTIONS:
- Extract all visible text and data
- Preserve structure and organization where possible
- Include tables, lists, and formatted content

RETURN: Complete readable content from this document."""
                        
                        # Handle file types based on Gemini API support
                        logging.info(f"STEP 1: Processing {file_name} with MIME type {mime_type}")
                        
                        # Check if this is a supported file type for Gemini API
                        if 'pdf' in mime_type or file_name.lower().endswith('.pdf'):
                            # PDF files - use chunked extraction for large files
                            logging.info(f"STEP 1: Using chunked PDF extraction for {file_name}")
                            
                            try:
                                # Import chunked extractor
                                from chunked_pdf_extractor import ChunkedPDFExtractor
                                
                                # Initialize extractor
                                extractor = ChunkedPDFExtractor()
                                
                                # Estimate PDF size and determine processing approach
                                pdf_info = extractor.estimate_pdf_size(binary_content)
                                logging.info(f"STEP 1: PDF analysis for {file_name}: {pdf_info}")
                                
                                if pdf_info["needs_chunking"]:
                                    logging.info(f"STEP 1: Large PDF detected, using chunked extraction for {file_name}")
                                    # Use chunked extraction
                                    import base64 as b64
                                    base64_data = f"data:application/pdf;base64,{b64.b64encode(binary_content).decode()}"
                                    chunk_result = extractor.extract_pdf_chunked(base64_data, file_name)
                                    
                                    if chunk_result["success"]:
                                        document_content = chunk_result["extracted_text"]
                                        logging.info(f"STEP 1: Chunked extraction successful for {file_name} - {chunk_result['successful_chunks']}/{chunk_result['chunks_processed']} chunks processed")
                                    else:
                                        logging.error(f"STEP 1: Chunked extraction failed for {file_name}: {chunk_result.get('error', 'Unknown error')}")
                                        document_content = f"[ERROR EXTRACTING PDF WITH CHUNKING: {file_name}]\nError: {chunk_result.get('error', 'Unknown error')}"
                                else:
                                    logging.info(f"STEP 1: Small PDF, using single-pass extraction for {file_name}")
                                    # Use single-pass extraction for small PDFs
                                    single_result = extractor.extract_single_pdf(binary_content, file_name)
                                    
                                    if single_result["success"]:
                                        document_content = single_result["extracted_text"]
                                        logging.info(f"STEP 1: Single-pass extraction successful for {file_name}")
                                    else:
                                        logging.error(f"STEP 1: Single-pass extraction failed for {file_name}: {single_result.get('error', 'Unknown error')}")
                                        document_content = f"[ERROR EXTRACTING PDF: {file_name}]\nError: {single_result.get('error', 'Unknown error')}"
                                
                                # Set content_response to None since we handled extraction directly
                                content_response = None
                                
                            except ImportError:
                                logging.warning("STEP 1: Chunked extractor not available, falling back to direct Gemini API")
                                # Fallback to original direct extraction
                                try:
                                    content_response = model.generate_content(
                                        [
                                            {
                                                "mime_type": mime_type,
                                                "data": binary_content
                                            },
                                            extraction_prompt
                                        ],
                                        generation_config=genai.GenerationConfig(
                                            max_output_tokens=65536,
                                            temperature=0.1
                                        ),
                                        request_options={"timeout": 300}
                                    )
                                    logging.info(f"STEP 1: Fallback extraction successful for {file_name}")
                                except Exception as e:
                                    logging.error(f"STEP 1: Fallback extraction failed for {file_name}: {e}")
                                    content_response = None
                            except Exception as e:
                                logging.error(f"STEP 1: Chunked extraction system error for {file_name}: {e}")
                                document_content = f"[ERROR IN CHUNKED EXTRACTION SYSTEM: {file_name}]\nError: {str(e)}"
                                content_response = None
                        elif ('word' in mime_type or 
                              'vnd.openxmlformats-officedocument.wordprocessingml' in mime_type or
                              'application/msword' in mime_type or
                              file_name.lower().endswith(('.docx', '.doc'))):
                            # Word document - use python-docx library
                            logging.info(f"STEP 1: Using python-docx library for Word extraction from {file_name}")
                            try:
                                import io
                                from docx import Document
                                
                                # Create document from binary content
                                doc_stream = io.BytesIO(binary_content)
                                doc = Document(doc_stream)
                                
                                # Extract all text content
                                text_content = []
                                for paragraph in doc.paragraphs:
                                    if paragraph.text.strip():
                                        text_content.append(paragraph.text.strip())
                                
                                # Extract text from tables
                                for table in doc.tables:
                                    for row in table.rows:
                                        row_text = []
                                        for cell in row.cells:
                                            if cell.text.strip():
                                                row_text.append(cell.text.strip())
                                        if row_text:
                                            text_content.append(" | ".join(row_text))
                                
                                document_content = "\n".join(text_content)
                                logging.info(f"STEP 1: Successfully extracted {len(document_content)} characters from Word document {file_name}")
                                
                            except Exception as e:
                                logging.error(f"STEP 1: Failed to extract content from Word document {file_name}: {e}")
                                document_content = f"[ERROR EXTRACTING WORD DOCUMENT: {file_name}]\nError: {str(e)}"
                        
                        elif ('excel' in mime_type or 
                              'spreadsheet' in mime_type or 
                              'vnd.ms-excel' in mime_type or 
                              'vnd.openxmlformats-officedocument.spreadsheetml' in mime_type or
                              file_name.lower().endswith(('.xlsx', '.xls'))):
                            # Excel file - use openpyxl/xlrd libraries
                            logging.info(f"STEP 1: Using openpyxl/xlrd libraries for Excel extraction from {file_name}")
                            try:
                                import io
                                import pandas as pd
                                
                                # Create Excel stream from binary content
                                excel_stream = io.BytesIO(binary_content)
                                
                                # Read all sheets
                                all_sheets = pd.read_excel(excel_stream, sheet_name=None, engine='openpyxl' if file_name.lower().endswith('.xlsx') else 'xlrd')
                                
                                text_content = []
                                total_content_size = 0
                                MAX_CONTENT_SIZE = 800000  # Increased to 800KB per document for better Excel coverage
                                
                                for sheet_name, df in all_sheets.items():
                                    text_content.append(f"=== SHEET: {sheet_name} ===")
                                    
                                    # For very large datasets, sample more intelligently  
                                    if len(df) > 500:  # If more than 500 rows, sample more extensively
                                        logging.info(f"Very large Excel sheet detected ({len(df)} rows), using intelligent sampling")
                                        # Take larger samples to capture more data patterns
                                        sample_df = pd.concat([
                                            df.head(150),  # First 150 rows (often headers + significant data)
                                            df.iloc[len(df)//4:len(df)//4+50],  # 50 rows from 1st quarter
                                            df.iloc[len(df)//2-25:len(df)//2+25],  # 50 middle rows
                                            df.iloc[3*len(df)//4:3*len(df)//4+50],  # 50 rows from 3rd quarter
                                            df.tail(100)   # Last 100 rows (often important summary data)
                                        ]).drop_duplicates()
                                        sheet_text = sample_df.to_string(index=False, na_rep='')
                                        sheet_text += f"\n\n[NOTE: This sheet has {len(df)} total rows. Showing comprehensive sample of {len(sample_df)} rows covering key data sections for complete analysis.]"
                                    elif len(df) > 200:  # Medium-large sheets
                                        logging.info(f"Large Excel sheet detected ({len(df)} rows), using expanded sampling")
                                        sample_df = pd.concat([
                                            df.head(100),  # First 100 rows
                                            df.iloc[len(df)//2-25:len(df)//2+25],  # 50 middle rows
                                            df.tail(50)   # Last 50 rows
                                        ]).drop_duplicates()
                                        sheet_text = sample_df.to_string(index=False, na_rep='')
                                        sheet_text += f"\n\n[NOTE: This sheet has {len(df)} total rows. Showing expanded sample of {len(sample_df)} rows for thorough analysis.]"
                                    else:
                                        # Convert full dataframe to string representation for smaller sheets
                                        sheet_text = df.to_string(index=False, na_rep='')
                                    
                                    # Check if adding this sheet would exceed size limit
                                    if total_content_size + len(sheet_text) > MAX_CONTENT_SIZE:
                                        # Truncate the sheet content
                                        remaining_space = MAX_CONTENT_SIZE - total_content_size
                                        if remaining_space > 1000:  # Only add if we have reasonable space
                                            sheet_text = sheet_text[:remaining_space] + "\n\n[CONTENT TRUNCATED - Sheet too large]"
                                            text_content.append(sheet_text)
                                            total_content_size += len(sheet_text)
                                        text_content.append(f"\n[REMAINING SHEETS SKIPPED - Document size limit reached]")
                                        break
                                    
                                    text_content.append(sheet_text)
                                    total_content_size += len(sheet_text)
                                
                                document_content = "\n\n".join(text_content)
                                logging.info(f"Excel extraction: Final content size {len(document_content)} chars (limit: {MAX_CONTENT_SIZE})")
                                logging.info(f"STEP 1: Successfully extracted {len(document_content)} characters from Excel document {file_name}")
                                
                            except Exception as e:
                                logging.error(f"STEP 1: Failed to extract content from Excel document {file_name}: {e}")
                                document_content = f"[ERROR EXTRACTING EXCEL DOCUMENT: {file_name}]\nError: {str(e)}"
                        
                        else:
                            # Other unsupported formats
                            logging.warning(f"STEP 1: Unsupported file format {file_name} ({mime_type})")
                            document_content = f"[UNSUPPORTED FILE FORMAT: {file_name}]\nFile type {mime_type} is not supported for content extraction."
                        
                        # Handle PDF response or use extracted content from libraries
                        if 'pdf' in mime_type or file_name.lower().endswith('.pdf'):
                            # Handle PDF response from Gemini API
                            if content_response and content_response.text:
                                document_content = content_response.text.strip()
                                logging.info(f"STEP 1: Successfully extracted {len(document_content)} characters from PDF {file_name}")
                            else:
                                logging.warning(f"STEP 1: No content extracted from PDF {file_name}")
                                document_content = "[Content extraction failed]"
                        
                        # Add extracted content to final text
                        extracted_content_text += f"\n\n=== DOCUMENT: {file_name} ===\n{document_content}"
                    
                except Exception as e:
                    logging.error(f"STEP 1: Failed to extract content from document {file_name}: {e}")
                    extracted_content_text += f"\n\n=== DOCUMENT: {file_name} ===\n[Content extraction error: {e}]"
                    continue
        
        logging.info(f"STEP 1 COMPLETE: Processed {processed_docs} documents, extracted total of {len(extracted_content_text)} characters from all documents")
        
        # Check total content size and truncate if needed to prevent timeouts
        MAX_TOTAL_CONTENT = 2500000  # Increased to 2.5MB total limit for very large Excel processing
        if len(extracted_content_text) > MAX_TOTAL_CONTENT:
            logging.warning(f"Total content size ({len(extracted_content_text)} chars) exceeds limit ({MAX_TOTAL_CONTENT}), truncating...")
            extracted_content_text = extracted_content_text[:MAX_TOTAL_CONTENT] + "\n\n[CONTENT TRUNCATED - Document set too large for single processing]"
        
        # Now proceed with data extraction using the extracted content
        final_prompt = full_prompt + f"\n\nEXTRACTED DOCUMENT CONTENT:\n{extracted_content_text}"
        
        # STEP 2: DATA EXTRACTION FROM CONTENT
        logging.info(f"=== STEP 2: DATA EXTRACTION ===")
        logging.info(f"Making data extraction call with {len(extracted_content_text)} characters of extracted content")
        
        # Retry logic for API overload situations in data extraction
        max_retries = 3
        retry_delay = 2  # Start with 2 seconds
        
        for attempt in range(max_retries):
            try:
                response = model.generate_content(
                    final_prompt,
                    generation_config=genai.GenerationConfig(
                        max_output_tokens=100000,  # Increased output limit for comprehensive Excel extraction
                        temperature=0.1,
                        response_mime_type="application/json"
                    ),
                    request_options={"timeout": None}  # Remove timeout constraints
                )
                break  # Success, exit retry loop
            except Exception as e:
                if "503" in str(e) and "overloaded" in str(e).lower() and attempt < max_retries - 1:
                    logging.warning(f"Data extraction API overloaded (attempt {attempt + 1}/{max_retries}), retrying in {retry_delay}s...")
                    import time
                    time.sleep(retry_delay)
                    retry_delay *= 2  # Exponential backoff
                    continue
                else:
                    raise e  # Re-raise if not overload error or final attempt
        
        if not response or not response.text:
            return ExtractionResult(success=False, error_message="No response from AI")
        
        # Extract token usage information
        input_token_count = None
        output_token_count = None
        
        # Check for different possible token usage structures in Gemini API response
        logging.info(f"RESPONSE DEBUG: Response type: {type(response)}")
        logging.info(f"RESPONSE DEBUG: Response attributes: {dir(response)}")
        
        if hasattr(response, 'usage_metadata') and response.usage_metadata:
            logging.info(f"TOKEN DEBUG: Found usage_metadata: {response.usage_metadata}")
            logging.info(f"TOKEN DEBUG: usage_metadata attributes: {dir(response.usage_metadata)}")
            input_token_count = getattr(response.usage_metadata, 'prompt_token_count', None)
            output_token_count = getattr(response.usage_metadata, 'candidates_token_count', None)
            
            # Try alternative attribute names if the primary ones don't exist
            if input_token_count is None:
                input_token_count = getattr(response.usage_metadata, 'input_token_count', None)
                input_token_count = getattr(response.usage_metadata, 'prompt_tokens', input_token_count)
            if output_token_count is None:
                output_token_count = getattr(response.usage_metadata, 'output_token_count', None) 
                output_token_count = getattr(response.usage_metadata, 'candidates_tokens', output_token_count)
                output_token_count = getattr(response.usage_metadata, 'completion_tokens', output_token_count)
            
            logging.info(f"TOKEN USAGE: Input tokens: {input_token_count}, Output tokens: {output_token_count}")
        else:
            logging.warning("TOKEN DEBUG: No usage_metadata found in response")
            
        # Fallback: try to find usage information in other response attributes
        if input_token_count is None and hasattr(response, '_result') and hasattr(response._result, 'usage_metadata'):
            usage = response._result.usage_metadata
            logging.info(f"TOKEN DEBUG: Found _result.usage_metadata: {usage}")
            input_token_count = getattr(usage, 'prompt_token_count', None)
            output_token_count = getattr(usage, 'candidates_token_count', None)
            logging.info(f"TOKEN USAGE (fallback): Input tokens: {input_token_count}, Output tokens: {output_token_count}")
        
        # Parse JSON response - handle markdown code blocks
        response_text = response.text.strip()
        
        # Log raw AI response for debugging
        logging.info(f"STEP 2: Raw AI response length: {len(response_text)}")
        logging.info(f"STEP 2: Raw AI response start: {response_text[:500]}...")
        if len(response_text) > 500:
            logging.info(f"STEP 2: Raw AI response end: ...{response_text[-500:]}")
            
        # Check for potential truncation indicators
        truncation_indicators = ['...', '...}', '"ai_reasoning": "Extracted from document analysis",']
        if any(indicator in response_text[-100:] for indicator in truncation_indicators):
            logging.warning("POTENTIAL RESPONSE TRUNCATION DETECTED - Response may be incomplete")
            
        # Count field_validations objects to detect truncation
        field_validation_count = response_text.count('"field_id":')
        logging.info(f"RECORD COUNT CHECK: Found {field_validation_count} field_validation objects in response")
        
        if field_validation_count > 100:
            logging.warning(f"AI returned {field_validation_count} records, exceeding 100-record limit")
        
        # Check for session-specific troubleshooting
        if "0db04e6a-006b-48af-b9bb-b1dc88edaae5" in str(session_name):
            logging.info("DEBUGGING SESSION 0db04e6a: Checking for collection extraction completeness")
            increase_rates_count = response_text.count('"collection_name": "Increase Rates"')
            logging.info(f"DEBUGGING SESSION 0db04e6a: Found {increase_rates_count} Increase Rates collection items")
            if increase_rates_count < 10:
                logging.warning(f"DEBUGGING SESSION 0db04e6a: Only found {increase_rates_count} Increase Rates items, expected 10!")
                logging.warning("DEBUGGING SESSION 0db04e6a: This may indicate response truncation or incomplete extraction")
            
        # Check if response appears to end abruptly
        if len(response_text) > 10000 and not response_text.strip().endswith((']}', '}')):
            logging.warning("RESPONSE APPEARS TRUNCATED - Does not end with expected JSON closing")
            
        # Check for 100-record limit compliance
        if field_validation_count > 100:
            logging.warning(f"Response contains {field_validation_count} records - will enforce 100-record limit")
        
        # Remove markdown code blocks if present
        if response_text.startswith("```json"):
            response_text = response_text[7:]  # Remove ```json
        if response_text.startswith("```"):
            response_text = response_text[3:]   # Remove ```
        if response_text.endswith("```"):
            response_text = response_text[:-3]  # Remove trailing ```
        
        response_text = response_text.strip()
        
        try:
            # If empty response, create default structure
            if not response_text:
                logging.warning("Empty response from AI, creating default structure")
                extracted_data = {"field_validations": []}
            else:
                # First attempt: direct JSON parsing
                try:
                    extracted_data = json.loads(response_text)
                except json.JSONDecodeError as e:
                    logging.warning(f"📋 Direct JSON parsing failed: {e}")
                    logging.warning("🔧 Attempting JSON repair for truncated response...")
                    
                    # Attempt to repair truncated JSON
                    repaired_json = repair_truncated_json(response_text)
                    if repaired_json:
                        try:
                            extracted_data = json.loads(repaired_json)
                            logging.info("✅ Successfully repaired and parsed truncated JSON response")
                            logging.info(f"🔢 Recovered {len(extracted_data.get('field_validations', []))} field validations from truncated response")
                        except json.JSONDecodeError:
                            logging.error("❌ JSON repair failed, creating minimal structure")
                            extracted_data = {"field_validations": []}
                    else:
                        logging.error("❌ Could not repair JSON, creating minimal structure")
                        extracted_data = {"field_validations": []}
                
            # Validate that we have the expected field_validations structure
            if "field_validations" not in extracted_data:
                logging.warning("AI response missing field_validations key, creating default structure")
                extracted_data = {"field_validations": []}
            
            # Enforce 100-record limit and ensure complete collection items
            field_validations = extracted_data.get('field_validations', [])
            if len(field_validations) > 100:
                logging.warning(f"AI returned {len(field_validations)} records, enforcing 100-record limit")
                
                # Find the last complete collection item within the limit
                truncated_validations = []
                collection_items = {}
                
                for i, validation in enumerate(field_validations[:100]):
                    truncated_validations.append(validation)
                    
                    # Track collection items to ensure completeness
                    if validation.get('validation_type') == 'collection_property':
                        collection_name = validation.get('collection_name', '')
                        record_index = validation.get('record_index', 0)
                        
                        if collection_name not in collection_items:
                            collection_items[collection_name] = {}
                        if record_index not in collection_items[collection_name]:
                            collection_items[collection_name][record_index] = []
                        
                        collection_items[collection_name][record_index].append(validation)
                
                # Check if we need to remove incomplete collection items at the boundary
                # Look ahead to see if there are more properties for the last collection items
                if len(field_validations) > 100:
                    # Check if the next validation(s) would complete a collection item
                    boundary_collections = {}
                    for validation in field_validations[100:]:
                        if validation.get('validation_type') == 'collection_property':
                            collection_name = validation.get('collection_name', '')
                            record_index = validation.get('record_index', 0)
                            key = f"{collection_name}_{record_index}"
                            if key not in boundary_collections:
                                boundary_collections[key] = []
                            boundary_collections[key].append(validation)
                    
                    # Remove incomplete collection items from the truncated list
                    final_validations = []
                    incomplete_collection_items = set()
                    
                    for validation in truncated_validations:
                        if validation.get('validation_type') == 'collection_property':
                            collection_name = validation.get('collection_name', '')
                            record_index = validation.get('record_index', 0)
                            key = f"{collection_name}_{record_index}"
                            
                            # If this collection item continues beyond the 100-record limit, mark as incomplete
                            if key in boundary_collections:
                                incomplete_collection_items.add(key)
                            
                            # Only include if this collection item is complete within the limit
                            if key not in incomplete_collection_items:
                                final_validations.append(validation)
                        else:
                            # Always include schema fields
                            final_validations.append(validation)
                    
                    extracted_data['field_validations'] = final_validations
                    logging.info(f"Trimmed to {len(final_validations)} complete records, removed {len(truncated_validations) - len(final_validations)} incomplete collection items")
                else:
                    extracted_data['field_validations'] = truncated_validations
                    logging.info(f"Trimmed to exactly 100 records")
                
            logging.info(f"STEP 1: Successfully extracted {len(extracted_data.get('field_validations', []))} field validations")
            return ExtractionResult(
                success=True, 
                extracted_data=extracted_data,
                extraction_prompt=final_prompt,
                ai_response=response.text,
                input_token_count=input_token_count,
                output_token_count=output_token_count
            )
            
        except json.JSONDecodeError as e:
            logging.error(f"Failed to parse AI response as JSON: {e}")
            logging.error(f"Cleaned response: {response_text}")
            return ExtractionResult(success=False, error_message=f"Invalid JSON response: {e}")
            
    except Exception as e:
        logging.error(f"STEP 1 extraction failed: {e}")
        return ExtractionResult(success=False, error_message=str(e))

# Validation function removed - validation now occurs only during extraction process

# Chain function removed - only extraction is needed, validation occurs during extraction

if __name__ == "__main__":
    import sys
    import json
    
    try:
        # Read input from stdin
        raw_input = sys.stdin.read()
        logging.info(f"RAW INPUT RECEIVED: {raw_input[:500]}...")
        input_data = json.loads(raw_input)
        
        # Log parsed data structure
        logging.info(f"PARSED INPUT KEYS: {list(input_data.keys())}")
        operation = input_data.get("step", input_data.get("operation", "extract"))
        
        if operation == "extract_text_only":
            # Extract text from documents without AI analysis
            documents = input_data.get("documents", [])
            
            if not documents:
                print(json.dumps({"success": False, "error": "No documents provided"}), file=sys.stderr)
                sys.exit(1)
            
            extracted_texts = []
            for doc in documents:
                try:
                    file_name = doc.get("file_name", "unknown")
                    file_content = doc.get("file_content", "")
                    mime_type = doc.get("mime_type", "")
                    
                    # Extract actual text based on file type
                    if file_content.startswith("data:"):
                        try:
                            if file_name.lower().endswith(('.xlsx', '.xls')) or 'spreadsheet' in mime_type:
                                # Handle Excel files using existing pandas/openpyxl processing
                                text_content = process_excel_file_content(file_content, file_name)
                                if not text_content or text_content.strip() == "":
                                    text_content = f"No data could be extracted from {file_name}"
                            elif file_name.lower().endswith('.pdf') or 'pdf' in mime_type:
                                # Handle PDF files using Gemini
                                logging.info(f"Extracting PDF content from {file_name}")
                                text_content = extract_pdf_with_gemini(file_content, file_name)
                                logging.info(f"PDF extraction result for {file_name}: {len(text_content) if text_content else 0} characters")
                                if not text_content or text_content.strip() == "":
                                    text_content = f"No text could be extracted from {file_name}"
                                    logging.warning(f"Empty content returned for PDF {file_name}")
                                else:
                                    # Log first 200 characters to verify content
                                    logging.info(f"PDF content preview: {text_content[:200]}...")
                            else:
                                # Handle other file types (DOCX, etc.)
                                text_content = extract_document_with_gemini(file_content, file_name, mime_type)
                                if not text_content or text_content.strip() == "":
                                    text_content = f"No text could be extracted from {file_name}"
                            
                            word_count = len(text_content.split()) if text_content else 0
                        except Exception as extract_error:
                            logging.error(f"Failed to extract content from {file_name}: {extract_error}")
                            text_content = f"Error extracting content from {file_name}: {str(extract_error)}"
                            word_count = 0
                    else:
                        # Handle plain text content
                        text_content = file_content[:1000]  # Truncate for safety
                        word_count = len(text_content.split())
                    
                    extracted_texts.append({
                        "file_name": file_name,
                        "text_content": text_content,
                        "word_count": word_count
                    })
                    
                except Exception as e:
                    logging.error(f"Error processing {doc.get('file_name', 'unknown')}: {e}")
                    extracted_texts.append({
                        "file_name": doc.get("file_name", "unknown"),
                        "text_content": f"Error processing file: {e}",
                        "word_count": 0
                    })
            
            # Log the final extracted texts for debugging
            logging.info(f"FINAL RESULT: {len(extracted_texts)} documents processed")
            for i, text_info in enumerate(extracted_texts):
                logging.info(f"Document {i+1}: {text_info['file_name']} - {text_info['word_count']} words, content length: {len(text_info.get('text_content', ''))}")
                if text_info.get('text_content'):
                    logging.info(f"Content preview: {text_info['text_content'][:200]}...")
                else:
                    logging.warning(f"Empty content for {text_info['file_name']}")
            
            print(json.dumps({
                "success": True,
                "extracted_texts": extracted_texts
            }))
            
        elif operation in ["extract", "extract_additional"]:
            # Extract from documents (regular or additional)
            documents = input_data.get("files", input_data.get("documents", []))  # Support both parameter names
            project_schema = input_data.get("project_schema", {})
            extraction_rules = input_data.get("extraction_rules", [])
            knowledge_documents = input_data.get("knowledge_documents", [])
            session_name = input_data.get("session_name", "contract")
            validated_data_context = input_data.get("validated_data_context", {})  # New context for subsequent uploads
            extraction_notes = input_data.get("extraction_notes", "")  # Special extraction instructions
            is_subsequent_upload = input_data.get("is_subsequent_upload", False)  # Flag for upload type
            
            # Call the extraction function with new parameters
            result = step1_extract_from_documents(documents, project_schema, extraction_rules, knowledge_documents, session_name, validated_data_context, extraction_notes, is_subsequent_upload)
            
            if result.success:
                print(json.dumps({
                    "success": True, 
                    "extracted_data": result.extracted_data, 
                    "field_validations": result.extracted_data.get("field_validations", []) if result.extracted_data else [],
                    "extraction_prompt": result.extraction_prompt,
                    "ai_response": result.ai_response,
                    "input_token_count": result.input_token_count,
                    "output_token_count": result.output_token_count
                }))
            else:
                print(json.dumps({"success": False, "error": result.error_message}), file=sys.stderr)
                sys.exit(1)
                
        else:
            print(json.dumps({"error": f"Unknown operation: {operation}"}), file=sys.stderr)
            sys.exit(1)
            
    except Exception as e:
        print(json.dumps({"error": str(e)}), file=sys.stderr)
        sys.exit(1)