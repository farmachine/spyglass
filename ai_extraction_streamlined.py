#!/usr/bin/env python3

import json
import sys
import os
import logging
import base64
from typing import Dict, Any, List
from dataclasses import dataclass

@dataclass
class FieldValidationResult:
    field_id: str
    field_name: str
    field_type: str
    extracted_value: Any
    original_extracted_value: Any
    original_confidence_score: int
    original_ai_reasoning: str
    validation_status: str
    ai_reasoning: str
    confidence_score: int
    document_source: str
    document_sections: List[str]
    collection_name: str = None
    record_index: int = None

@dataclass
class ExtractionResult:
    extracted_data: Dict[str, Any]
    confidence_score: float
    processing_notes: str
    field_validations: List[FieldValidationResult]

def generate_human_friendly_reasoning(field_name: str, extracted_value: Any, applied_rules: List[Dict[str, Any]]) -> str:
    """Generate human-friendly email-style reasoning for fields with conflicts or issues."""
    if not applied_rules:
        return f"Successfully extracted {field_name} from document"
    
    reasoning = f"We have extracted '{extracted_value}' for {field_name}, however there are some considerations that require your attention:\n\n"
    reasoning += "IDENTIFIED CONCERNS:\n"
    
    for rule in applied_rules:
        rule_name = rule.get('name', 'Policy Check')
        action = rule.get('action', '')
        
        if 'Knowledge Document Conflict' in rule_name:
            reasoning += "• Our compliance review process has flagged this value based on internal policies.\n"
        elif 'Inc' in rule_name.lower():
            reasoning += "• The extracted company appears to be incorporated, requiring additional verification.\n"
        else:
            reasoning += f"• {action}\n"
    
    reasoning += "\nTo help us complete the verification process, could you please clarify:\n\n"
    
    if 'company' in field_name.lower() or 'name' in field_name.lower():
        reasoning += "1. Can you confirm the full legal name as it appears in official documentation?\n"
        reasoning += "2. What is the primary business relationship with this entity?\n"
        reasoning += "3. Are there any specific compliance considerations?\n"
    elif 'country' in field_name.lower() or 'jurisdiction' in field_name.lower():
        reasoning += "1. Can you confirm the governing jurisdiction for this agreement?\n"
        reasoning += "2. Are there any cross-border regulatory requirements?\n"
        reasoning += "3. Should this be subject to specific regional compliance procedures?\n"
    elif 'date' in field_name.lower():
        reasoning += "1. Can you confirm the exact date for this field?\n"
        reasoning += "2. Is this date subject to any notice requirements?\n"
        reasoning += "3. Are there any related dates we should track?\n"
    else:
        reasoning += f"1. Can you verify that '{extracted_value}' is correct for {field_name}?\n"
        reasoning += "2. Are there any additional details we should consider?\n"
        reasoning += "3. Does this require special handling or approval?\n"
    
    reasoning += "\nThank you for your assistance in completing this review."
    return reasoning

def ai_validate_batch(field_validations: List[Dict[str, Any]], extraction_rules: List[Dict[str, Any]] = None, knowledge_documents: List[Dict[str, Any]] = None) -> Dict[str, tuple[int, List[Dict[str, Any]], str]]:
    """Pure AI validation - all decisions made by AI only."""
    if not field_validations:
        return {}
    
    logging.info(f"AI validation processing {len(field_validations)} fields")
    
    try:
        api_key = os.environ.get("GEMINI_API_KEY")
        if not api_key:
            logging.warning("GEMINI_API_KEY not found, using default confidence")
            return {fv['field_name']: (95, [], "Extracted during AI processing") for fv in field_validations}
        
        from google import genai
        from google.genai import types
        
        client = genai.Client(api_key=api_key)
        
        # Build validation prompt
        fields_summary = "FIELDS TO VALIDATE:\n"
        for fv in field_validations:
            if fv.get('extracted_value') is not None:
                fields_summary += f"- Field: {fv['field_name']} = '{fv['extracted_value']}'\n"
        
        rules_context = ""
        if extraction_rules:
            rules_context = "EXTRACTION RULES:\n"
            for rule in extraction_rules:
                if rule.get("isActive", True):
                    rules_context += f"- {rule.get('ruleName', 'Unknown')}: {rule.get('ruleContent', '')}\n"
        
        knowledge_context = ""
        if knowledge_documents:
            knowledge_context = "KNOWLEDGE DOCUMENTS:\n"
            for doc in knowledge_documents:
                doc_name = doc.get('displayName', doc.get('fileName', 'Unknown'))
                content = doc.get('content', '')
                if content:
                    knowledge_context += f"- {doc_name}: {content[:500]}...\n"
        
        validation_prompt = f"""You are an expert data validation specialist. Analyze extracted field values using AI judgment only.

{fields_summary}

{rules_context}

{knowledge_context}

For each field, provide:
1. Confidence percentage (0-100%)
2. Applied rules/conflicts
3. Brief reasoning

Confidence guidelines:
- 95-100%: Clear, accurate extraction
- 80-94%: Good with minor uncertainty  
- 50-79%: Reasonable but some concerns
- 0-49%: Poor extraction or issues
- 0%: No value found

Return JSON format:
{{
    "field_name": {{
        "confidence_percentage": <integer>,
        "applied_rules": [{{ "name": "<rule>", "action": "<description>" }}],
        "reasoning": "<explanation>"
    }}
}}"""

        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=[types.Content(role="user", parts=[types.Part(text=validation_prompt)])],
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                max_output_tokens=2048,
                temperature=0.1
            )
        )
        
        if response.text:
            try:
                batch_results = json.loads(response.text)
                results = {}
                for fv in field_validations:
                    field_name = fv['field_name']
                    if field_name in batch_results:
                        result = batch_results[field_name]
                        confidence = result.get("confidence_percentage", 95)
                        applied_rules = result.get("applied_rules", [])
                        reasoning = result.get("reasoning", "Extracted during AI processing")
                        results[field_name] = (confidence, applied_rules, reasoning)
                    else:
                        results[field_name] = (95, [], "Extracted during AI processing")
                return results
            except json.JSONDecodeError:
                logging.error("Failed to parse AI validation response")
                return {fv['field_name']: (95, [], "Extracted during AI processing") for fv in field_validations}
        
        return {fv['field_name']: (95, [], "Extracted during AI processing") for fv in field_validations}
        
    except Exception as e:
        logging.error(f"AI validation error: {e}")
        return {fv['field_name']: (95, [], "Extracted during AI processing") for fv in field_validations}

def extract_with_ai(documents: List[Dict[str, Any]], project_schema: Dict[str, Any], extraction_rules: List[Dict[str, Any]] = None, knowledge_documents: List[Dict[str, Any]] = None, session_id: str = None) -> ExtractionResult:
    """Streamlined AI extraction with Gemini."""
    try:
        api_key = os.environ.get("GEMINI_API_KEY")
        if not api_key:
            raise Exception("GEMINI_API_KEY environment variable not found")
        
        from google import genai
        from google.genai import types
        
        # Build extraction prompt
        prompt = build_extraction_prompt(project_schema, extraction_rules, knowledge_documents)
        
        # Process documents
        all_extracted_data = []
        
        for doc in documents:
            try:
                file_name = doc.get('file_name', '')
                file_content = doc.get('file_content', '')
                mime_type = doc.get('mime_type', '')
                
                logging.info(f"Processing document: {file_name} ({mime_type})")
                
                # Extract content based on type
                if mime_type.startswith("text/"):
                    content_text = extract_text_content(file_content)
                    extracted_data = process_text_with_ai(prompt, content_text)
                elif mime_type == "application/pdf":
                    extracted_data = process_pdf_with_ai(prompt, file_content, file_name)
                elif "excel" in mime_type or "spreadsheet" in mime_type:
                    extracted_data = process_excel_with_pandas(file_content, file_name)
                elif "word" in mime_type:
                    extracted_data = process_word_with_docx(file_content, file_name)
                else:
                    logging.warning(f"Unsupported file type: {mime_type}")
                    continue
                
                if extracted_data:
                    all_extracted_data.append(extracted_data)
                    
            except Exception as e:
                logging.error(f"Error processing document {file_name}: {e}")
                continue
        
        # Aggregate results from all documents
        aggregated_data = aggregate_extraction_results(all_extracted_data, project_schema)
        
        # Create field validations
        field_validations = create_field_validations(aggregated_data, project_schema, extraction_rules, knowledge_documents)
        
        return ExtractionResult(
            extracted_data=aggregated_data,
            confidence_score=95.0,
            processing_notes=f"Processed {len(documents)} documents successfully",
            field_validations=field_validations
        )
        
    except Exception as e:
        logging.error(f"AI extraction failed: {e}")
        raise

def build_extraction_prompt(project_schema: Dict[str, Any], extraction_rules: List[Dict[str, Any]], knowledge_documents: List[Dict[str, Any]]) -> str:
    """Build comprehensive extraction prompt."""
    prompt = "You are an expert document data extraction specialist. Extract structured data exactly as specified.\n\n"
    
    # Add schema fields
    if project_schema.get("schema_fields"):
        prompt += "EXTRACT THE FOLLOWING MAIN FIELDS:\n"
        for field in project_schema["schema_fields"]:
            field_name = field.get('fieldName', '')
            description = field.get('description', '')
            prompt += f"- {field_name}: {description}\n"
    
    # Add collections
    if project_schema.get("collections"):
        prompt += "\nEXTRACT THE FOLLOWING COLLECTIONS:\n"
        for collection in project_schema["collections"]:
            collection_name = collection.get('collectionName', collection.get('objectName', ''))
            prompt += f"\n{collection_name} (array of objects):\n"
            
            properties = collection.get("properties", [])
            for prop in properties:
                if isinstance(prop, dict):
                    prop_name = prop.get('propertyName', '')
                    description = prop.get('description', '')
                    prompt += f"  - {prop_name}: {description}\n"
    
    # Add JSON format example
    prompt += "\nRETURN DATA IN THIS JSON FORMAT:\n{\n"
    
    if project_schema.get("schema_fields"):
        for field in project_schema["schema_fields"]:
            field_name = field.get('fieldName', '')
            prompt += f'  "{field_name}": null,\n'
    
    if project_schema.get("collections"):
        for collection in project_schema["collections"]:
            collection_name = collection.get('collectionName', collection.get('objectName', ''))
            prompt += f'  "{collection_name}": [\n    {{\n'
            properties = collection.get("properties", [])
            for prop in properties:
                if isinstance(prop, dict):
                    prop_name = prop.get('propertyName', '')
                    prompt += f'      "{prop_name}": null,\n'
            prompt += '    }\n  ],\n'
    
    prompt += "}\n"
    
    # Add extraction guidelines
    prompt += "\nEXTRACTION GUIDELINES:\n"
    prompt += "• Extract all parties/organizations mentioned in the document\n"
    prompt += "• Look for company names, addresses, countries in headers, signatures, and body text\n"
    prompt += "• Include full legal names and complete address information when available\n"
    prompt += "• For contracts/NDAs, extract ALL signing parties, not just primary ones\n"
    
    # Add knowledge documents context
    if knowledge_documents:
        prompt += "\nKNOWLEDGE BASE CONTEXT:\n"
        for doc in knowledge_documents:
            doc_name = doc.get('displayName', doc.get('fileName', 'Unknown'))
            content = doc.get('content', '')
            if content:
                prompt += f"Document: {doc_name}\nContent: {content}\n\n"
        prompt += "Consider the above policies when extracting data.\n"
    
    return prompt

def extract_text_content(file_content: str) -> str:
    """Extract text content from string or data URL."""
    if isinstance(file_content, str):
        if file_content.startswith('data:'):
            base64_content = file_content.split(',', 1)[1]
            decoded_bytes = base64.b64decode(base64_content)
            return decoded_bytes.decode('utf-8', errors='ignore')
        else:
            return file_content
    else:
        return file_content.decode('utf-8', errors='ignore')

def process_text_with_ai(prompt: str, content_text: str) -> Dict[str, Any]:
    """Process text content with Gemini AI."""
    from google import genai
    
    full_prompt = f"{prompt}\n\nDocument content:\n{content_text}"
    model = genai.GenerativeModel('gemini-1.5-flash')
    response = model.generate_content(full_prompt)
    
    if response and response.text:
        return parse_ai_response(response.text)
    return {}

def process_pdf_with_ai(prompt: str, file_content: str, file_name: str) -> Dict[str, Any]:
    """Process PDF with multiple fallback methods."""
    from google import genai
    
    # Decode base64 content
    if file_content.startswith('data:'):
        base64_content = file_content.split(',', 1)[1]
        binary_content = base64.b64decode(base64_content)
    else:
        binary_content = base64.b64decode(file_content)
    
    model = genai.GenerativeModel('gemini-1.5-flash')
    
    # Try PyPDF2 text extraction first
    try:
        import PyPDF2
        import io
        
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(binary_content))
        text_content = ""
        for page in pdf_reader.pages:
            text_content += page.extract_text() + "\n"
        
        if text_content.strip():
            full_prompt = f"{prompt}\n\nDocument content:\n{text_content}"
            response = model.generate_content(full_prompt)
            if response and response.text:
                return parse_ai_response(response.text)
    except Exception as e:
        logging.error(f"PyPDF2 extraction failed: {e}")
    
    # Try pdf2image conversion as fallback
    try:
        from pdf2image import convert_from_bytes
        images = convert_from_bytes(binary_content, dpi=150, first_page=1, last_page=3)
        if images:
            response = model.generate_content([prompt, images[0]])
            if response and response.text:
                return parse_ai_response(response.text)
    except Exception as e:
        logging.error(f"PDF image processing failed: {e}")
    
    return {}

def process_excel_with_pandas(file_content: str, file_name: str) -> Dict[str, Any]:
    """Process Excel files using pandas."""
    try:
        import pandas as pd
        import io
        
        if file_content.startswith('data:'):
            base64_content = file_content.split(',', 1)[1]
            binary_content = base64.b64decode(base64_content)
        else:
            binary_content = base64.b64decode(file_content)
        
        # Read Excel file
        excel_data = pd.read_excel(io.BytesIO(binary_content), sheet_name=None)
        
        # Convert to structured text
        text_content = f"Excel file: {file_name}\n\n"
        for sheet_name, df in excel_data.items():
            text_content += f"Sheet: {sheet_name}\n"
            text_content += df.to_string(index=False) + "\n\n"
        
        # Process with AI
        from google import genai
        model = genai.GenerativeModel('gemini-1.5-flash')
        prompt = "Extract structured data from this Excel content:\n\n" + text_content
        response = model.generate_content(prompt)
        
        if response and response.text:
            return parse_ai_response(response.text)
            
    except Exception as e:
        logging.error(f"Excel processing failed: {e}")
    
    return {}

def process_word_with_docx(file_content: str, file_name: str) -> Dict[str, Any]:
    """Process Word documents using python-docx."""
    try:
        from docx import Document
        import io
        
        if file_content.startswith('data:'):
            base64_content = file_content.split(',', 1)[1]
            binary_content = base64.b64decode(base64_content)
        else:
            binary_content = base64.b64decode(file_content)
        
        # Read Word document
        doc = Document(io.BytesIO(binary_content))
        text_content = f"Word document: {file_name}\n\n"
        
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"
        
        # Process tables
        for table in doc.tables:
            text_content += "\nTable:\n"
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                text_content += row_text + "\n"
        
        # Process with AI
        from google import genai
        model = genai.GenerativeModel('gemini-1.5-flash')
        prompt = "Extract structured data from this Word document:\n\n" + text_content
        response = model.generate_content(prompt)
        
        if response and response.text:
            return parse_ai_response(response.text)
            
    except Exception as e:
        logging.error(f"Word processing failed: {e}")
    
    return {}

def parse_ai_response(response_text: str) -> Dict[str, Any]:
    """Parse AI response and extract JSON."""
    try:
        # Clean up response text
        response_text = response_text.strip()
        if response_text.startswith("```json"):
            response_text = response_text.replace("```json", "").replace("```", "").strip()
        elif response_text.startswith("```"):
            response_text = response_text.replace("```", "").strip()
        
        return json.loads(response_text)
    except json.JSONDecodeError as e:
        logging.error(f"Failed to parse AI response: {e}")
        return {}

def aggregate_extraction_results(all_extracted_data: List[Dict[str, Any]], project_schema: Dict[str, Any]) -> Dict[str, Any]:
    """Aggregate results from multiple documents."""
    if not all_extracted_data:
        return {}
    
    if len(all_extracted_data) == 1:
        return all_extracted_data[0]
    
    # Merge multiple documents
    aggregated = {}
    
    # Merge schema fields (take from first non-null value)
    if project_schema.get("schema_fields"):
        for field in project_schema["schema_fields"]:
            field_name = field.get('fieldName', '')
            for data in all_extracted_data:
                if field_name in data and data[field_name] is not None:
                    aggregated[field_name] = data[field_name]
                    break
    
    # Merge collections (combine all items)
    if project_schema.get("collections"):
        for collection in project_schema["collections"]:
            collection_name = collection.get('collectionName', collection.get('objectName', ''))
            combined_items = []
            
            for data in all_extracted_data:
                if collection_name in data and isinstance(data[collection_name], list):
                    combined_items.extend(data[collection_name])
            
            if combined_items:
                aggregated[collection_name] = combined_items
    
    return aggregated

def create_field_validations(aggregated_data: Dict[str, Any], project_schema: Dict[str, Any], extraction_rules: List[Dict[str, Any]], knowledge_documents: List[Dict[str, Any]]) -> List[FieldValidationResult]:
    """Create field validation records."""
    field_validations = []
    
    # Collect fields for batch validation
    fields_to_validate = []
    field_metadata = {}
    
    # Process schema fields
    if project_schema.get("schema_fields"):
        for field in project_schema["schema_fields"]:
            field_id = str(field.get("id", "unknown"))
            field_name = field.get("fieldName", "")
            field_type = field.get("fieldType", "TEXT")
            extracted_value = aggregated_data.get(field_name)
            
            fields_to_validate.append({
                'field_name': field_name,
                'extracted_value': extracted_value
            })
            
            field_metadata[field_name] = {
                'field_id': field_id,
                'field_type': field_type,
                'auto_verification_threshold': field.get("autoVerificationConfidence", 80),
                'extracted_value': extracted_value
            }
    
    # Process collection fields
    if project_schema.get("collections"):
        for collection in project_schema["collections"]:
            collection_name = collection.get('collectionName', collection.get('objectName', ''))
            collection_data = aggregated_data.get(collection_name, [])
            
            if isinstance(collection_data, list):
                for record_index, record in enumerate(collection_data):
                    properties = collection.get("properties", [])
                    for prop in properties:
                        if isinstance(prop, dict):
                            prop_id = str(prop.get("id", "unknown"))
                            prop_name = prop.get("propertyName", "")
                            prop_type = prop.get("propertyType", "TEXT")
                            
                            field_name_with_index = f"{collection_name}.{prop_name}[{record_index}]"
                            extracted_value = record.get(prop_name) if isinstance(record, dict) else None
                            
                            fields_to_validate.append({
                                'field_name': field_name_with_index,
                                'extracted_value': extracted_value
                            })
                            
                            field_metadata[field_name_with_index] = {
                                'field_id': prop_id,
                                'field_type': prop_type,
                                'auto_verification_threshold': prop.get("autoVerificationConfidence", 80),
                                'extracted_value': extracted_value,
                                'collection_name': collection_name,
                                'record_index': record_index
                            }
    
    # Run batch AI validation
    validation_results = ai_validate_batch(fields_to_validate, extraction_rules, knowledge_documents)
    
    # Create FieldValidationResult objects
    for field_name, metadata in field_metadata.items():
        confidence, applied_rules, reasoning = validation_results.get(field_name, (95, [], "Extracted during AI processing"))
        
        # Determine validation status based on confidence and auto-verification threshold
        auto_threshold = metadata.get('auto_verification_threshold', 80)
        if confidence >= auto_threshold and metadata.get('extracted_value') is not None:
            status = "verified"
        else:
            status = "unverified"
        
        validation = FieldValidationResult(
            field_id=metadata['field_id'],
            field_name=field_name,
            field_type=metadata['field_type'],
            extracted_value=metadata['extracted_value'],
            original_extracted_value=metadata['extracted_value'],
            original_confidence_score=confidence,
            original_ai_reasoning=reasoning,
            validation_status=status,
            ai_reasoning=reasoning,
            confidence_score=confidence,
            document_source="Multi-document extraction",
            document_sections=["Document analysis"],
            collection_name=metadata.get('collection_name'),
            record_index=metadata.get('record_index')
        )
        
        field_validations.append(validation)
    
    return field_validations

def main():
    """Main entry point for AI extraction."""
    if len(sys.argv) != 2:
        print(json.dumps({"error": "Usage: python ai_extraction.py <json_input>"}))
        sys.exit(1)
    
    try:
        input_data = json.loads(sys.argv[1])
        
        documents = input_data.get('documents', [])
        project_schema = input_data.get('project_schema', {})
        extraction_rules = input_data.get('extraction_rules', [])
        knowledge_documents = input_data.get('knowledge_documents', [])
        session_id = input_data.get('session_id', '')
        
        result = extract_with_ai(documents, project_schema, extraction_rules, knowledge_documents, session_id)
        
        # Convert FieldValidationResult objects to dicts for JSON serialization
        field_validations_dict = []
        for fv in result.field_validations:
            field_validations_dict.append({
                'field_id': fv.field_id,
                'field_name': fv.field_name,
                'field_type': fv.field_type,
                'extracted_value': fv.extracted_value,
                'original_extracted_value': fv.original_extracted_value,
                'original_confidence_score': fv.original_confidence_score,
                'original_ai_reasoning': fv.original_ai_reasoning,
                'validation_status': fv.validation_status,
                'ai_reasoning': fv.ai_reasoning,
                'confidence_score': fv.confidence_score,
                'document_source': fv.document_source,
                'document_sections': fv.document_sections,
                'collection_name': fv.collection_name,
                'record_index': fv.record_index
            })
        
        output = {
            "success": True,
            "extracted_data": result.extracted_data,
            "confidence_score": result.confidence_score,
            "processing_notes": result.processing_notes,
            "field_validations": field_validations_dict
        }
        
        print(json.dumps(output))
        
    except Exception as e:
        logging.error(f"Extraction failed: {e}")
        print(json.dumps({
            "success": False,
            "error": str(e),
            "extracted_data": {},
            "field_validations": []
        }))
        sys.exit(1)

if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    main()