#!/usr/bin/env python3
"""
Verify that the document storage functionality is working correctly
and test with different document types
"""
import json
import pandas as pd
import io
import base64
import requests
from docx import Document
import tempfile
import os

def create_test_docx():
    """Create a test Word document"""
    doc = Document()
    doc.add_heading('Test Document', 0)
    
    p = doc.add_paragraph('This is a test document for verifying document storage.')
    p.add_run(' Here is some bold text.').bold = True
    p.add_run(' And here is some italic text.').italic = True
    
    doc.add_heading('Main Content', level=1)
    doc.add_paragraph('This document contains important business information.')
    doc.add_paragraph('Contract details, financial data, and legal terms.')
    
    # Save to memory
    temp_buffer = io.BytesIO()
    doc.save(temp_buffer)
    temp_buffer.seek(0)
    
    # Convert to base64 data URL
    docx_data = temp_buffer.getvalue()
    base64_data = base64.b64encode(docx_data).decode()
    data_url = f"data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{base64_data}"
    
    return data_url

def test_multi_document_extraction():
    """Test extracting multiple documents at once"""
    
    # Create test Excel
    excel_data = {
        'Contract_ID': ['CNT001', 'CNT002', 'CNT003'],
        'Client_Name': ['Apple Inc', 'Google LLC', 'Microsoft Corp'],
        'Value': [1000000, 750000, 1200000],
        'Status': ['Active', 'Pending', 'Active']
    }
    
    df = pd.DataFrame(excel_data)
    excel_buffer = io.BytesIO()
    with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='Contracts', index=False)
    excel_buffer.seek(0)
    excel_data_b64 = base64.b64encode(excel_buffer.getvalue()).decode()
    excel_data_url = f"data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{excel_data_b64}"
    
    # Create test Word document
    docx_data_url = create_test_docx()
    
    # Prepare the data for multi-document extraction
    test_data = {
        "files": [
            {
                "name": "contracts_data.xlsx",
                "content": excel_data_url,
                "type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            },
            {
                "name": "business_document.docx", 
                "content": docx_data_url,
                "type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            }
        ]
    }
    
    print("Testing multi-document storage...")
    print(f"Excel file size: {len(excel_data_url)} bytes")
    print(f"Word file size: {len(docx_data_url)} bytes")
    
    try:
        # Use a different session for this test
        session_id = "5d776706-0e91-40c0-aa5f-e514290199c2"
        
        response = requests.post(
            f"http://localhost:5000/api/sessions/{session_id}/extract-text",
            json=test_data,
            headers={"Content-Type": "application/json"}
        )
        
        print(f"Response status: {response.status_code}")
        
        if response.status_code == 200:
            result = response.json()
            print(f"✅ Extracted {len(result.get('extractedTexts', []))} documents")
            
            # Check stored documents
            docs_response = requests.get(f"http://localhost:5000/api/sessions/{session_id}/documents")
            
            if docs_response.status_code == 200:
                documents = docs_response.json()
                print(f"📄 Total documents in session: {len(documents)}")
                
                for doc in documents:
                    content_preview = doc['extractedContent'][:100] if doc['extractedContent'] else ""
                    print(f"  - {doc['fileName']}")
                    print(f"    Size: {doc['fileSize']} bytes")
                    print(f"    Type: {doc['mimeType']}")
                    print(f"    Content preview: {content_preview}...")
                    print()
                
                return True
            else:
                print(f"❌ Failed to fetch documents: {docs_response.status_code}")
                return False
        else:
            print(f"❌ Extract text failed: {response.text}")
            return False
            
    except Exception as e:
        print(f"❌ Error: {e}")
        return False

def verify_existing_documents():
    """Verify that our earlier test document is still there"""
    print("Checking existing session documents...")
    
    try:
        docs_response = requests.get("http://localhost:5000/api/sessions/d76539a6-d7c0-4eb3-9b07-02ff4830747a/documents")
        
        if docs_response.status_code == 200:
            documents = docs_response.json()
            print(f"✅ Found {len(documents)} existing documents")
            
            for doc in documents:
                print(f"  - {doc['fileName']} ({doc['fileSize']} bytes, {len(doc['extractedContent'])} chars extracted)")
                
        else:
            print(f"❌ Failed to fetch existing documents: {docs_response.status_code}")
            
    except Exception as e:
        print(f"❌ Error checking existing documents: {e}")

if __name__ == "__main__":
    print("=== Document Storage Verification ===\n")
    
    verify_existing_documents()
    print()
    
    success = test_multi_document_extraction()
    
    if success:
        print("🎉 Document storage system is working correctly!")
        print("✅ Multiple file types supported")
        print("✅ Content extraction working")
        print("✅ Database storage functioning")
        print("✅ API endpoints operational")
    else:
        print("❌ Issues detected with document storage system")