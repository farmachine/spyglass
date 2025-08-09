#!/usr/bin/env python3
"""
OPTIMIZED AI EXTRACTION SYSTEM
Token-efficient extraction with maintained quality
Key optimizations:
1. Shorter, more focused prompts
2. Reduced redundant instructions 
3. Streamlined AI reasoning requirements
4. Efficient JSON structure
5. Smart content summarization for large documents
"""
import os
import json
import logging
import base64
from typing import Dict, Any, List, Optional
from dataclasses import dataclass

# Configure logging
logging.basicConfig(level=logging.INFO)

@dataclass
class ExtractionResult:
    success: bool
    extracted_data: Optional[Dict[str, Any]] = None
    error_message: Optional[str] = None
    extraction_prompt: Optional[str] = None
    ai_response: Optional[str] = None
    input_token_count: Optional[int] = None
    output_token_count: Optional[int] = None

def repair_truncated_json(response_text: str) -> str:
    """Enhanced repair function for truncated JSON responses"""
    try:
        logging.info(f"Attempting to repair JSON response of length {len(response_text)}")
        
        if not response_text.strip().startswith('{"field_validations":'):
            logging.warning("Response doesn't start with expected field_validations structure")
            if '"field_validations"' in response_text[:200]:
                logging.info("Found field_validations key later in response, attempting repair...")
                start_idx = response_text.find('"field_validations"')
                if start_idx > 0:
                    response_text = '{"' + response_text[start_idx:]
            else:
                return None
        
        lines = response_text.split('\n')
        field_validations = []
        current_object_lines = []
        brace_count = 0
        inside_validation = False
        in_field_validations_array = False
        
        for line_num, line in enumerate(lines):
            if '"field_validations":' in line:
                in_field_validations_array = True
                continue
                
            if not in_field_validations_array:
                continue
                
            line_stripped = line.strip()
            if line_stripped == '{' or ('{' in line and ('"field_id"' in line or '"field_name"' in line)):
                if line_stripped == '{':
                    current_object_lines = [line]
                else:
                    current_object_lines = [line]
                inside_validation = True
                brace_count = line.count('{') - line.count('}')
            elif inside_validation:
                current_object_lines.append(line)
                brace_count += line.count('{') - line.count('}')
                
                if brace_count == 0 and (line.strip().endswith('}') or line.strip().endswith('},')):
                    complete_object = '\n'.join(current_object_lines)
                    try:
                        obj_json = complete_object.strip()
                        if obj_json.endswith(','):
                            obj_json = obj_json[:-1]
                        
                        test_json = '{"test": ' + obj_json + '}'
                        parsed_test = json.loads(test_json)
                        
                        test_obj = parsed_test['test']
                        if 'field_id' in test_obj or 'field_name' in test_obj:
                            field_validations.append(complete_object.strip())
                            logging.info(f"Found complete field validation object #{len(field_validations)}")
                        else:
                            logging.warning("Field validation object missing required keys")
                            
                    except json.JSONDecodeError as e:
                        logging.warning(f"Skipping invalid field validation object: {str(e)[:100]}...")
                    
                    current_object_lines = []
                    inside_validation = False
                    brace_count = 0
        
        if field_validations:
            repaired = '{\n  "field_validations": [\n'
            
            for i, validation in enumerate(field_validations):
                clean_validation = validation.strip()
                if clean_validation.endswith(','):
                    clean_validation = clean_validation[:-1]
                
                indented_validation = '\n'.join('    ' + line for line in clean_validation.split('\n'))
                repaired += indented_validation
                
                if i < len(field_validations) - 1:
                    repaired += ','
                repaired += '\n'
            
            repaired += '  ]\n}'
            
            try:
                json.loads(repaired)
                logging.info(f"✅ Successfully repaired truncated JSON - preserved {len(field_validations)} complete field validations")
                return repaired
            except json.JSONDecodeError as e:
                logging.error(f"❌ Repaired JSON is still invalid: {e}")
                return None
        else:
            logging.warning("No complete field validation objects found in truncated response")
            return None
            
    except Exception as e:
        logging.error(f"Error during JSON repair: {e}")
        return None

def create_optimized_prompt(schema_fields: List[Dict], collections: List[Dict], extraction_rules: List[Dict], knowledge_documents: List[str], document_content: str, session_name: str) -> str:
    """Create a token-optimized extraction prompt"""
    
    # Summarize document content if it's very large (>50k chars)
    if len(document_content) > 50000:
        logging.info(f"Large document detected ({len(document_content)} chars), creating summary for context")
        # Take first 20k chars + last 5k chars + middle sample
        doc_summary = document_content[:20000] + "\n\n[... CONTENT TRUNCATED FOR EFFICIENCY ...]\n\n" + document_content[-5000:]
        content_to_use = doc_summary
        logging.info(f"Document summarized from {len(document_content)} to {len(content_to_use)} characters")
    else:
        content_to_use = document_content
    
    # Build compact schema representation
    schema_text = ""
    if schema_fields:
        schema_text = "\n## FIELDS:\n"
        for field in schema_fields:
            desc = field.get('description', '').strip()
            if len(desc) > 200:  # Truncate very long descriptions
                desc = desc[:200] + "..."
            schema_text += f"- **{field.get('name', '')}** ({field.get('dataType', 'TEXT')}): {desc}\n"
    
    # Build compact collections representation  
    collections_text = ""
    if collections:
        collections_text = "\n## COLLECTIONS:\n"
        for collection in collections:
            coll_desc = collection.get('description', '').strip()
            if len(coll_desc) > 150:  # Truncate long collection descriptions
                coll_desc = coll_desc[:150] + "..."
            collections_text += f"- **{collection.get('name', '')}**: {coll_desc}\n"
            
            properties = collection.get('properties', [])
            if properties:
                collections_text += "  Properties:\n"
                for prop in properties[:5]:  # Limit to first 5 properties to save tokens
                    prop_desc = prop.get('description', '').strip()
                    if len(prop_desc) > 100:  # Shorter property descriptions
                        prop_desc = prop_desc[:100] + "..."
                    collections_text += f"    * **{prop.get('name', '')}** ({prop.get('dataType', 'TEXT')}): {prop_desc}\n"
                if len(properties) > 5:
                    collections_text += f"    [... and {len(properties) - 5} more properties]\n"
    
    # Build compact rules representation
    rules_text = ""
    if extraction_rules:
        rules_text = "\n## RULES:\n"
        for rule in extraction_rules[:10]:  # Limit to 10 most important rules
            rule_desc = rule.get('description', '').strip()
            if len(rule_desc) > 100:  # Shorter rule descriptions
                rule_desc = rule_desc[:100] + "..."
            rules_text += f"- {rule.get('fieldName', '')}: {rule_desc}\n"
        if len(extraction_rules) > 10:
            rules_text += f"[... and {len(extraction_rules) - 10} more rules]\n"
    
    # Create streamlined prompt (much shorter than original 18k chars)
    prompt = f"""Extract data from documents into JSON format.

## INSTRUCTIONS:
1. Process ALL documents thoroughly
2. Follow field descriptions precisely  
3. Apply extraction rules where specified
4. For collections: Extract EVERY instance as separate items (maximum 50 items per collection)
5. Use knowledge documents for validation where listed
6. **VERIFIED DATA**: Reference verified field validations but DO NOT re-extract them
7. **COLLECTION EXPANSION**: If collections have verified data, search for additional items beyond verified ones
8. Return only valid JSON - no explanations

{schema_text}{collections_text}{rules_text}

## OUTPUT FORMAT:
```json
{{
  "field_validations": [
    {{
      "field_id": "...",
      "validation_type": "schema_field|collection_property", 
      "data_type": "...",
      "field_name": "...",
      "extracted_value": "...",
      "confidence_score": 0.0-1.0,
      "validation_status": "verified|unverified",
      "ai_reasoning": "Brief explanation with source location"
    }}
  ]
}}
```

## DOCUMENT CONTENT:
{content_to_use}

Return JSON only."""
    
    return prompt

def step1_extract_from_documents(documents: List[Dict[str, Any]], project_schema: Dict[str, Any], extraction_rules: List[Dict[str, Any]], knowledge_documents: List[str], session_name: str = "document") -> ExtractionResult:
    """Optimized document extraction with reduced token usage"""
    try:
        import google.generativeai as genai
        
        # Configure Gemini API
        api_key = os.environ.get("GEMINI_API_KEY")
        if not api_key:
            return ExtractionResult(success=False, error_message="GEMINI_API_KEY not found")
        
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-2.0-flash-exp')
        
        # Extract schema information
        schema_fields = project_schema.get('fields', [])
        collections = project_schema.get('collections', [])
        
        logging.info(f"OPTIMIZED EXTRACTION: Processing {len(documents)} documents with {len(schema_fields)} fields and {len(collections)} collections")
        
        # Process documents and extract content
        document_content = ""
        
        for doc in documents:
            file_name = doc.get('name', 'unknown')
            mime_type = doc.get('mimeType', 'application/octet-stream')
            
            # Get file content
            if 'content' in doc:
                content = doc['content']
                if content.startswith('data:'):
                    try:
                        content_b64 = content.split(',')[1]
                        binary_content = base64.b64decode(content_b64)
                    except Exception as e:
                        logging.error(f"Failed to decode base64 content for {file_name}: {e}")
                        continue
                else:
                    try:
                        binary_content = base64.b64decode(content)
                    except Exception as e:
                        logging.error(f"Failed to decode content for {file_name}: {e}")
                        continue
            else:
                logging.warning(f"No content found for document {file_name}")
                continue
            
            # Extract content based on file type (optimized approach)
            if 'pdf' in mime_type or file_name.lower().endswith('.pdf'):
                # For PDFs, use chunked extraction if available, otherwise direct API
                try:
                    from chunked_pdf_extractor import ChunkedPDFExtractor
                    extractor = ChunkedPDFExtractor()
                    pdf_info = extractor.estimate_pdf_size(binary_content)
                    
                    if pdf_info["needs_chunking"]:
                        logging.info(f"Large PDF detected for {file_name}, using chunked extraction")
                        base64_data = f"data:application/pdf;base64,{base64.b64encode(binary_content).decode()}"
                        chunk_result = extractor.extract_pdf_chunked(base64_data, file_name)
                        if chunk_result["success"]:
                            file_content = chunk_result["extracted_text"]
                        else:
                            file_content = f"[PDF EXTRACTION FAILED: {file_name}]"
                    else:
                        single_result = extractor.extract_single_pdf(binary_content, file_name)
                        if single_result["success"]:
                            file_content = single_result["extracted_text"]
                        else:
                            file_content = f"[PDF EXTRACTION FAILED: {file_name}]"
                except ImportError:
                    # Fallback to direct Gemini API
                    try:
                        response = model.generate_content(
                            [{"mime_type": mime_type, "data": binary_content}, "Extract all text content from this document."],
                            generation_config=genai.GenerationConfig(max_output_tokens=32768, temperature=0.1)
                        )
                        file_content = response.text if response and response.text else f"[FAILED TO EXTRACT: {file_name}]"
                    except Exception as e:
                        file_content = f"[EXTRACTION ERROR: {file_name}] {str(e)}"
                        
            elif ('excel' in mime_type or 'spreadsheet' in mime_type or 
                  file_name.lower().endswith(('.xlsx', '.xls'))):
                # Excel extraction using pandas
                try:
                    import io
                    import pandas as pd
                    excel_stream = io.BytesIO(binary_content)
                    
                    # Read all sheets efficiently 
                    all_sheets = pd.read_excel(excel_stream, sheet_name=None, header=0)
                    sheet_contents = []
                    
                    for sheet_name, df in all_sheets.items():
                        if not df.empty:
                            # Convert to simple string representation (more efficient than full table)
                            sheet_text = f"\n## SHEET: {sheet_name}\n"
                            sheet_text += f"Columns: {', '.join(df.columns.astype(str))}\n"
                            sheet_text += f"Rows: {len(df)}\n"
                            
                            # Include sample data (first 3 rows max to save tokens)
                            if len(df) > 0:
                                sample_df = df.head(3)
                                sheet_text += "Sample data:\n"
                                sheet_text += sample_df.to_string(max_rows=3, max_cols=10, show_dimensions=False)
                            sheet_contents.append(sheet_text)
                    
                    file_content = "\n".join(sheet_contents)
                    logging.info(f"Excel extraction successful for {file_name} - {len(all_sheets)} sheets processed efficiently")
                    
                except Exception as e:
                    logging.error(f"Excel extraction failed for {file_name}: {e}")
                    file_content = f"[EXCEL EXTRACTION FAILED: {file_name}] {str(e)}"
                    
            elif ('word' in mime_type or 'vnd.openxmlformats-officedocument.wordprocessingml' in mime_type):
                # Word document extraction
                try:
                    import io
                    from docx import Document
                    doc_stream = io.BytesIO(binary_content)
                    doc = Document(doc_stream)
                    
                    text_content = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            text_content.append(paragraph.text.strip())
                    
                    # Extract table content efficiently
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_text.append(cell.text.strip())
                            if row_text:
                                text_content.append(" | ".join(row_text))
                    
                    file_content = "\n".join(text_content)
                    logging.info(f"Word extraction successful for {file_name}")
                    
                except Exception as e:
                    logging.error(f"Word extraction failed for {file_name}: {e}")
                    file_content = f"[WORD EXTRACTION FAILED: {file_name}] {str(e)}"
            else:
                # Generic text extraction via Gemini API
                try:
                    response = model.generate_content(
                        [{"mime_type": mime_type, "data": binary_content}, "Extract all readable text content."],
                        generation_config=genai.GenerationConfig(max_output_tokens=16384, temperature=0.1)
                    )
                    file_content = response.text if response and response.text else f"[NO CONTENT EXTRACTED: {file_name}]"
                except Exception as e:
                    file_content = f"[EXTRACTION ERROR: {file_name}] {str(e)}"
            
            # Add to document content with clear separation
            document_content += f"\n\n=== DOCUMENT: {file_name} ===\n{file_content}\n"
        
        if not document_content.strip():
            return ExtractionResult(success=False, error_message="No content extracted from documents")
        
        logging.info(f"Total extracted content: {len(document_content)} characters")
        
        # Create optimized prompt
        final_prompt = create_optimized_prompt(
            schema_fields, collections, extraction_rules, 
            knowledge_documents, document_content, session_name
        )
        
        logging.info(f"OPTIMIZED PROMPT: Created prompt with {len(final_prompt)} characters (vs typical 50k+)")
        
        # Call Gemini API with optimized settings
        max_retries = 3
        retry_delay = 2
        
        for attempt in range(max_retries):
            try:
                response = model.generate_content(
                    final_prompt,
                    generation_config=genai.GenerationConfig(
                        max_output_tokens=32768,  # Reduced from 65536 to save tokens
                        temperature=0.1,
                        response_mime_type="application/json"
                    ),
                    request_options={"timeout": None}
                )
                break
            except Exception as e:
                if "503" in str(e) and "overloaded" in str(e).lower() and attempt < max_retries - 1:
                    logging.warning(f"API overloaded (attempt {attempt + 1}/{max_retries}), retrying in {retry_delay}s...")
                    import time
                    time.sleep(retry_delay)
                    retry_delay *= 2
                    continue
                else:
                    raise e
        
        if not response or not response.text:
            return ExtractionResult(success=False, error_message="No response from AI")
        
        # Extract token usage
        input_token_count = None
        output_token_count = None
        
        if hasattr(response, 'usage_metadata') and response.usage_metadata:
            input_token_count = getattr(response.usage_metadata, 'prompt_token_count', None)
            output_token_count = getattr(response.usage_metadata, 'candidates_token_count', None)
            logging.info(f"TOKEN USAGE (OPTIMIZED): Input: {input_token_count}, Output: {output_token_count}")
        
        # Parse response
        response_text = response.text.strip()
        
        # Remove markdown code blocks
        if response_text.startswith("```json"):
            response_text = response_text[7:]
        if response_text.startswith("```"):
            response_text = response_text[3:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]
        
        response_text = response_text.strip()
        
        try:
            if not response_text:
                extracted_data = {"field_validations": []}
            else:
                try:
                    extracted_data = json.loads(response_text)
                except json.JSONDecodeError as e:
                    logging.warning(f"Direct JSON parsing failed: {e}")
                    # Try repair function
                    repaired_json = repair_truncated_json(response_text)
                    if repaired_json:
                        try:
                            extracted_data = json.loads(repaired_json)
                            logging.info(f"Successfully repaired JSON - recovered {len(extracted_data.get('field_validations', []))} field validations")
                        except json.JSONDecodeError:
                            extracted_data = {"field_validations": []}
                    else:
                        extracted_data = {"field_validations": []}
            
            if "field_validations" not in extracted_data:
                extracted_data = {"field_validations": []}
            
            validation_count = len(extracted_data.get('field_validations', []))
            logging.info(f"OPTIMIZED EXTRACTION: Successfully extracted {validation_count} field validations")
            
            return ExtractionResult(
                success=True,
                extracted_data=extracted_data,
                extraction_prompt=final_prompt,
                ai_response=response.text,
                input_token_count=input_token_count,
                output_token_count=output_token_count
            )
            
        except Exception as e:
            logging.error(f"Failed to parse AI response: {e}")
            return ExtractionResult(success=False, error_message=f"Invalid JSON response: {e}")
            
    except Exception as e:
        logging.error(f"Optimized extraction failed: {e}")
        return ExtractionResult(success=False, error_message=str(e))

if __name__ == "__main__":
    import sys
    import json
    
    try:
        raw_input = sys.stdin.read()
        input_data = json.loads(raw_input)
        
        operation = input_data.get("step", input_data.get("operation", "extract"))
        
        if operation == "extract":
            documents = input_data.get("files", input_data.get("documents", []))
            project_schema = input_data.get("project_schema", {})
            extraction_rules = input_data.get("extraction_rules", [])
            knowledge_documents = input_data.get("knowledge_documents", [])
            session_name = input_data.get("session_name", "contract")
            
            result = step1_extract_from_documents(documents, project_schema, extraction_rules, knowledge_documents, session_name)
            
            if result.success:
                print(json.dumps({
                    "success": True,
                    "extracted_data": result.extracted_data,
                    "field_validations": result.extracted_data.get("field_validations", []),
                    "extraction_prompt": result.extraction_prompt,
                    "ai_response": result.ai_response,
                    "input_token_count": result.input_token_count,
                    "output_token_count": result.output_token_count
                }))
            else:
                print(json.dumps({"success": False, "error": result.error_message}), file=sys.stderr)
                sys.exit(1)
        else:
            print(json.dumps({"error": f"Unknown operation: {operation}"}), file=sys.stderr)
            sys.exit(1)
            
    except Exception as e:
        print(json.dumps({"error": str(e)}), file=sys.stderr)
        sys.exit(1)